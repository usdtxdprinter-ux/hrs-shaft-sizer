"""
MES Exhaust Shaft Sizing Calculator — Streamlit Chatbot
=========================================================
A chatbot-style application for sizing fire-rated exhaust shafts
in high-rise buildings using the LF Systems MES constant pressure system.

Products: DEF (Dryer Exhaust Fan), DBF (Dryer Booster Fan), L150/L152 controllers
Website:  https://www.lfsystems.net
System:   MES (Multi-Family Exhaust Shaft)

Engineering Basis:
  - Darcy-Weisbach friction loss: Δpf = f*(L/Dh)*ρ*(V/1096.2)²
  - Colebrook friction factor for turbulent flow
  - Huebscher equivalent diameter for rectangular ducts
  - ASHRAE 2009 Duct Design Chapter 21 fitting loss coefficients
  - Subduct area deductions: 4"→15 sq.in., 6"→31.5 sq.in., 8"→54 sq.in.

Deploy:   pip install streamlit pandas plotly
          streamlit run mes_shaft_sizer.py
"""

import streamlit as st
import math
import pandas as pd
import json
import io
import os
import tempfile

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
AIR_DENSITY = 0.075          # lb/ft³ at standard conditions
ROUGHNESS   = 0.0003         # ft — galvanized steel, medium-smooth
KIN_VISC    = 1.63e-4        # ft²/s — air at ~70 °F

SUBDUCT_AREA = {4: 15.0, 6: 31.5, 8: 54.0}  # sq.in. removed per penetration

ROUND_SIZES = [8,10,12,14,16,18,20,22,24,26,28,30,32,34,36,38,40,42,44,46,48]
RECT_SIDES  = [6,8,10,12,14,16,18,20,22,24,26,28,30,32,34,36,38,40,42,44,46,48]

# Fitting loss coefficients (ASHRAE basis)
K_ELBOW_90      = 0.60
K_ELBOW_45      = 0.30
K_TEE_BRANCH    = 0.50
K_EXIT          = 1.00
K_ENTRY_BELL    = 0.03
K_ENTRY_ABRUPT  = 0.50

# ─────────────────────────────────────────────
# DEF FAN CURVE DATA (from DEF_Fan_Curves.xlsx)
# Each fan: list of (CFM, SP in.WC) points
# ─────────────────────────────────────────────
DEF_FAN_CURVES = {
    "DEF004": {
        "cfm_sp": [(540,0),(490,0.25),(430,0.50),(350,0.75),(240,1.00)],
        "voltage": "120V/1ph", "hp": "1/2", "motor": "EC",
        "amps": "6.25", "weight": "35 lbs",
        "inlet": '8"', "impeller": "BI", "rpm": "1950",
    },
    "DEF008": {
        "cfm_sp": [(970,0),(890,0.25),(840,0.50),(780,0.75),(680,1.00),(540,1.25),(440,1.50),(270,1.75)],
        "voltage": "120V/1ph", "hp": "1/2", "motor": "EC",
        "amps": "6.25", "weight": "40 lbs",
        "inlet": '10"', "impeller": "BI", "rpm": "1950",
    },
    "DEF015": {
        "cfm_sp": [(1860,0),(1780,0.25),(1700,0.50),(1610,0.75),(1520,1.00),(1410,1.25),(1280,1.50),(1140,1.75),(990,2.00)],
        "voltage": "120V/1ph", "hp": "1/2", "motor": "EC",
        "amps": "6.25", "weight": "55 lbs",
        "inlet": '12"', "impeller": "BI", "rpm": "1950",
    },
    "DEF025": {
        "cfm_sp": [(2480,0),(2400,0.25),(2320,0.50),(2230,0.75),(2140,1.00),(2040,1.25),(1930,1.50),(1790,1.75),(1630,2.00)],
        "voltage": "120V/1ph", "hp": "1", "motor": "EC",
        "amps": "12.9", "weight": "80 lbs",
        "inlet": '14"', "impeller": "BI", "rpm": "1950",
    },
    "DEF035": {
        "cfm_sp": [(4100,0),(3940,0.25),(3770,0.50),(3610,0.75),(3460,1.00),(3300,1.25),(3120,1.50),(2900,1.75),(2630,2.00)],
        "voltage": "208/480V/3ph", "hp": "3", "motor": "ID",
        "amps": "9.8/4.9", "weight": "275 lbs",
        "inlet": '18"', "impeller": "BC", "rpm": "1740",
    },
    "DEF050": {
        "cfm_sp": [(5850,0),(5660,0.25),(5450,0.50),(5300,0.75),(5090,1.00),(4890,1.25),(4680,1.50),(4460,1.75),(4230,2.00)],
        "voltage": "208/480V/3ph", "hp": "5", "motor": "ID",
        "amps": "12.6/6.4", "weight": "380 lbs",
        "inlet": '20"', "impeller": "BC", "rpm": "1740",
    },
}

FAN_ORDER = ["DEF004", "DEF008", "DEF015", "DEF025", "DEF035", "DEF050"]


def fan_max_cfm_at_sp(fan_name: str, sp: float) -> float:
    """Interpolate fan curve to find CFM at a given static pressure."""
    pts = DEF_FAN_CURVES[fan_name]["cfm_sp"]
    # pts are (cfm, sp) pairs with SP ascending, CFM descending
    sps = [p[1] for p in pts]
    cfms = [p[0] for p in pts]
    if sp <= sps[0]:
        return cfms[0]
    if sp >= sps[-1]:
        return cfms[-1]
    for i in range(len(sps) - 1):
        if sps[i] <= sp <= sps[i + 1]:
            frac = (sp - sps[i]) / (sps[i + 1] - sps[i])
            return cfms[i] + frac * (cfms[i + 1] - cfms[i])
    return cfms[-1]


def fan_sp_at_cfm(fan_name: str, cfm: float) -> float:
    """Interpolate fan curve to find SP at a given CFM (inverse lookup)."""
    pts = DEF_FAN_CURVES[fan_name]["cfm_sp"]
    cfms = [p[0] for p in pts]  # descending
    sps = [p[1] for p in pts]   # ascending
    if cfm >= cfms[0]:
        return sps[0]
    if cfm <= cfms[-1]:
        return sps[-1]
    for i in range(len(cfms) - 1):
        if cfms[i] >= cfm >= cfms[i + 1]:
            frac = (cfms[i] - cfm) / (cfms[i] - cfms[i + 1])
            return sps[i] + frac * (sps[i + 1] - sps[i])
    return sps[-1]


def select_fan(design_cfm: float, system_sp: float) -> dict:
    """
    Select the best DEF fan for the application.
    If a single DEF050 isn't enough, recommend multiple DEF050s in parallel.
    Returns dict with fan selection details.
    """
    # Try each fan in order (smallest to largest)
    for name in FAN_ORDER:
        available_cfm = fan_max_cfm_at_sp(name, system_sp)
        if available_cfm >= design_cfm:
            # This fan can handle it
            op_sp = fan_sp_at_cfm(name, design_cfm)
            return {
                "model": name,
                "quantity": 1,
                "parallel": False,
                "design_cfm": design_cfm,
                "system_sp": round(system_sp, 4),
                "available_cfm": round(available_cfm, 0),
                "operating_sp": round(op_sp, 4),
                "margin_pct": round((available_cfm - design_cfm) / design_cfm * 100, 1),
                "specs": DEF_FAN_CURVES[name],
            }

    # If no single fan works, use multiple DEF050s in parallel
    # In parallel, each fan handles design_cfm / N at the same SP
    def050_max = fan_max_cfm_at_sp("DEF050", system_sp)
    num_fans = math.ceil(design_cfm / def050_max)
    cfm_per_fan = design_cfm / num_fans
    op_sp = fan_sp_at_cfm("DEF050", cfm_per_fan)

    return {
        "model": "DEF050",
        "quantity": num_fans,
        "parallel": True,
        "design_cfm": design_cfm,
        "cfm_per_fan": round(cfm_per_fan, 0),
        "system_sp": round(system_sp, 4),
        "available_cfm": round(def050_max * num_fans, 0),
        "operating_sp": round(op_sp, 4),
        "margin_pct": round((def050_max * num_fans - design_cfm) / design_cfm * 100, 1),
        "specs": DEF_FAN_CURVES["DEF050"],
    }


def select_controller(floors: int) -> dict:
    """Select L150-H for ≤7 stories, L152-M for >7 stories."""
    if floors <= 7:
        return {
            "model": "L150.H",
            "name": "L150 Constant Pressure Controller",
            "system": "HRS (High-Rise System)",
            "accessories": "LP5, SLT, DP",
            "reason": f"{floors} stories (≤7 → L150-H)",
            "listings": "UL508, CSA C22.2 No 14-18",
        }
    else:
        return {
            "model": "L152.M",
            "name": "L152 Constant Pressure Controller",
            "system": "MES (Multi-Story Exhaust System)",
            "accessories": "(2) LP5, (2) SLT, (2) DP",
            "reason": f"{floors} stories (>7 → L152 for stack effect mitigation)",
            "listings": "UL508, UL864, CSA C22.2 No 14-18",
        }


def compute_system_curve(params: dict, best: dict, n_points: int = 20) -> list:
    """
    Compute system curve: ΔP vs CFM from 0 to 120% of design CFM.
    System ΔP ∝ CFM² (approximately, for turbulent flow).
    Returns list of (cfm, sp) tuples.
    """
    design_cfm = best["design_cfm"]
    system_sp = best["dp_total"]
    points = []
    for i in range(n_points + 1):
        frac = i / n_points * 1.2  # 0 to 120%
        cfm = design_cfm * frac
        # System resistance follows square law: SP = k * CFM²
        sp = system_sp * (frac ** 2) if frac > 0 else 0
        points.append((round(cfm, 0), round(sp, 4)))
    return points


def generate_fan_system_chart(fan_sel: dict, system_curve: list) -> bytes:
    """Generate fan curve + system curve chart as PNG bytes using matplotlib."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker

    # Brand font
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.sans-serif'] = ['Roboto', 'DejaVu Sans', 'Arial']

    fig, ax = plt.subplots(figsize=(10, 6))
    fig.patch.set_facecolor('#fafafa')

    # Fan curve
    fan_name = fan_sel["model"]
    pts = DEF_FAN_CURVES[fan_name]["cfm_sp"]
    fan_cfms = [p[0] for p in pts]
    fan_sps = [p[1] for p in pts]

    if fan_sel["parallel"] and fan_sel["quantity"] > 1:
        # For parallel fans, multiply CFM by quantity
        fan_cfms = [c * fan_sel["quantity"] for c in fan_cfms]
        label = f'{fan_sel["quantity"]}x {fan_name} (parallel)'
    else:
        label = fan_name

    ax.plot(fan_cfms, fan_sps, '-o', color='#234699', linewidth=2, markersize=5, label=f'{label} Fan Curve', zorder=3)

    # System curve
    sys_cfms = [p[0] for p in system_curve]
    sys_sps = [p[1] for p in system_curve]
    ax.plot(sys_cfms, sys_sps, '-', color='#b11f33', linewidth=2, label='System Curve', zorder=3)

    # Operating point
    op_cfm = fan_sel["design_cfm"]
    op_sp = fan_sel["system_sp"]
    ax.plot(op_cfm, op_sp, '*', color='#2a3853', markersize=18, label=f'Operating Point ({op_cfm:.0f} CFM, {op_sp:.2f}" WC)',
            zorder=5, markeredgecolor='#101820', markeredgewidth=0.5)

    ax.set_xlabel('Airflow (CFM)', fontsize=12, fontweight='bold', color='#2a3853')
    ax.set_ylabel('Static Pressure (in. WC)', fontsize=12, fontweight='bold', color='#2a3853')
    ax.set_title(f'Fan Curve vs System Curve — {label}', fontsize=14, fontweight='bold', color='#101820')
    ax.legend(loc='upper right', fontsize=10)
    ax.grid(True, alpha=0.2, color='#97999b')
    ax.set_xlim(left=0)
    ax.set_ylim(bottom=0)
    ax.set_facecolor('#fafafa')

    # LF Systems branding
    ax.text(0.01, 0.01, 'LF Systems by RM Manifold — lfsystems.net', transform=ax.transAxes,
            fontsize=8, color='#97999b', ha='left', va='bottom')

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# DATA SHEET PATHS (ZIP archives with JPEG pages)
# ─────────────────────────────────────────────
def _resolve_datasheet_path(filename: str) -> str:
    """Find a data sheet file in common locations."""
    candidates = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), filename),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "datasheets", filename),
        os.path.join("/mnt/project", filename),
    ]
    for c in candidates:
        if os.path.exists(c):
            return c
    return ""

DATASHEET_PATHS = {
    "DEF":  _resolve_datasheet_path("DEF_Data_Sheet_1025_1_1.pdf"),
    "L150": _resolve_datasheet_path("L150_Data_Sheet_0325_1_1.pdf"),
    "L152": _resolve_datasheet_path("L152_Data_Sheet_0126_1.pdf"),
}


def extract_datasheet_images(sheet_key: str) -> list:
    """
    Extract JPEG page images from a data sheet ZIP archive.
    Returns a list of (jpeg_bytes, width, height) tuples in page order.
    """
    import zipfile as _zf
    path = DATASHEET_PATHS.get(sheet_key, "")
    if not path or not os.path.exists(path):
        return []
    try:
        images = []
        with _zf.ZipFile(path, "r") as z:
            manifest = json.loads(z.read("manifest.json"))
            for page_info in sorted(manifest["pages"], key=lambda p: p["page_number"]):
                img_path = page_info["image"]["path"]
                dims = page_info["image"]["dimensions"]
                img_bytes = z.read(img_path)
                images.append((img_bytes, dims["width"], dims["height"]))
        return images
    except Exception:
        return []


def build_gamma_report_text(ss, best, fan_sel, ctrl) -> str:
    """Build the full Gamma input text dynamically from real calculation results."""
    from datetime import date

    project_name = ss.get('project_name', 'Untitled Project') if hasattr(ss, 'get') else getattr(ss, 'project_name', 'Untitled Project')
    project_location = ss.get('project_location', '') if hasattr(ss, 'get') else getattr(ss, 'project_location', '')
    operator_email = ss.get('operator_email', '') if hasattr(ss, 'get') else getattr(ss, 'operator_email', '')
    exhaust_type = ss.get('exhaust_type', '') if hasattr(ss, 'get') else getattr(ss, 'exhaust_type', '')
    floors = ss.get('floors', 0) if hasattr(ss, 'get') else getattr(ss, 'floors', 0)
    floor_data = ss.get('floor_data', []) if hasattr(ss, 'get') else getattr(ss, 'floor_data', [])
    has_subducts = ss.get('has_subducts', True) if hasattr(ss, 'get') else getattr(ss, 'has_subducts', True)
    diversity_pct = ss.get('diversity_pct', 100) if hasattr(ss, 'get') else getattr(ss, 'diversity_pct', 100)
    floor_height = ss.get('floor_height', 0) if hasattr(ss, 'get') else getattr(ss, 'floor_height', 0)
    duct_after_last = ss.get('duct_after_last', 0) if hasattr(ss, 'get') else getattr(ss, 'duct_after_last', 0)
    has_offset = ss.get('has_offset', False) if hasattr(ss, 'get') else getattr(ss, 'has_offset', False)
    max_delta_p = ss.get('max_delta_p', 0.25) if hasattr(ss, 'get') else getattr(ss, 'max_delta_p', 0.25)

    total_pens = sum(f.get("penetrations", 1) for f in floor_data)
    sub_size = floor_data[0].get("subduct_size", 0) if floor_data else 0
    cfm_per = floor_data[0].get("cfm_per_pen", 0) if floor_data else 0
    total_cfm_raw = sum(f.get("penetrations", 1) * f.get("cfm_per_pen", 0) for f in floor_data)
    design_cfm = best.get("design_cfm", 0)
    shaft_h = best.get("shaft_height", floors * floor_height)

    label = best.get("label", "")
    gross_area = best.get("shaft_area", 0)
    eff_area = best.get("eff_area", 0)
    dh = best.get("dh", 0)
    max_vel = best.get("max_vel", 0)
    dp_total = best.get("dp_total", 0)
    dp_shaft = best.get("dp_shaft", 0)
    dp_after = best.get("dp_after", 0)
    dp_offset = best.get("dp_offset", 0)
    dp_exit = best.get("dp_exit", 0)
    pass_fail = "PASS" if best.get("pass", False) else "FAIL"
    max_floor_dp = best.get("max_floor_dp", 0)

    floor_balance = best.get("floor_balance", [])

    qty_label = f'{fan_sel["quantity"]}x ' if fan_sel["quantity"] > 1 else ''
    fan_model = fan_sel["model"]

    ctrl_model = ctrl.get("model", "")
    ctrl_system = ctrl.get("system", "")
    ctrl_accessories = ctrl.get("accessories", "")
    ctrl_listings = ctrl.get("listings", "")
    ctrl_reason = ctrl.get("reason", "")

    today_str = date.today().strftime("%B %Y")

    fb_rows = ""
    show_floors = floor_balance[:20] if len(floor_balance) > 20 else floor_balance
    for fb in show_floors:
        pos = "Bottom" if fb["floor"] == 1 else ("Top" if fb["floor"] == floors else "")
        fb_rows += f'| {fb["floor"]}{" (" + pos + ")" if pos else ""} | {fb["cumul_cfm"]:,.0f} | {fb["velocity"]:,.0f} | {fb["section_dp"]:.4f} | {fb["accum_dp"]:.4f} |\n'

    if "L152" in ctrl_model:
        ctrl_page_title = "L152 — Constant Pressure Controller"
        ctrl_text = (
            "The L152 is a constant pressure controller used for the Multi-Story Exhaust System (MES). "
            "Two differential pressure transducers are installed above and below the neutral pressure plane "
            "to mitigate the seasonal stack effect. UL 864 listed for fire alarm integration.\n\n"
            "| Parameter | Value |\n|---|---|\n"
            "| Power Input | 85-264V AC |\n| Consumption | 0.7A@115VAC / 0.42A@230VAC |\n"
            "| Relay Output | 30-250V, 0.65A |\n| Control Signal | 0-10VDC |\n"
            "| Dimensions | 6\"W x 7\"H x 2-3/8\"D |\n\n"
            "**ACCESSORIES:** (2) LP5 Pressure Transducers, (2) SLT Silicone Tubing, (2) DP Duct Probes"
        )
    else:
        ctrl_page_title = "L150 — Constant Pressure Controller"
        ctrl_text = (
            "The L150 maintains a constant negative pressure in a high-rise exhaust shaft. "
            "Designed with EC-Flow Technology for rapid pressure response.\n\n"
            "| Parameter | Value |\n|---|---|\n"
            "| Power Input | 24V AC/DC |\n| Consumption | 112 mA |\n"
            "| Relay Output | 30-250V, 0.65A |\n| Control Signal | 0-10VDC |\n"
            "| Dimensions | 6\"W x 7\"H x 2-3/8\"D |\n\n"
            "**ACCESSORIES:** LP5 Pressure Transducer, SLT Silicone Tubing, DP Duct Probe"
        )

    subduct_row = ""
    if has_subducts and sub_size > 0:
        sub_deduct = {4: 15, 6: 31.5, 8: 54}.get(sub_size, 0)
        pens_per = floor_data[0].get("penetrations", 1) if floor_data else 1
        subduct_row = f"| Subduct Deductions ({pens_per} x {sub_size}\") | {sub_deduct * pens_per:.0f} sq. in. |\n"

    margin_pct = fan_sel.get("margin_pct", "N/A")

    text = f"""## PAGE 1 — COVER PAGE

# MES EXHAUST SHAFT SIZING REPORT

### Multi-Family Shaft Constant Pressure System

**Project:** {project_name}
**Location:** {project_location}
**Exhaust Type:** {exhaust_type}
**Prepared by:** {operator_email}
**Date:** {today_str}

---

**SYSTEM AT A GLANCE**

| | |
|---|---|
| Building Floors | {floors} |
| Design CFM | {design_cfm:,.0f} CFM |
| Recommended Shaft | {label} |
| Selected Fan | {qty_label}{fan_model} |
| Controller | {ctrl_model} |
| System Status | {pass_fail} |
| Total System dP | {dp_total:.3f} in. WC |

---

LF Systems by RM Manifold | 100 S Sylvania Ave, Fort Worth, TX 76111 | 817-393-4029 | lfsystems.net

---

## PAGE 2 — SYSTEM SUMMARY AND DESIGN INPUTS

### System Design Parameters

| Parameter | Value |
|---|---|
| Exhaust Type | {exhaust_type} |
| Number of Floors | {floors} |
| Penetrations per Floor | {floor_data[0].get("penetrations", 1) if floor_data else "N/A"} |
| Total Penetrations | {total_pens} |
{"| Subduct Size | " + str(sub_size) + '" diameter |' if has_subducts and sub_size > 0 else "| Subducts | None (wall openings) |"}
| CFM per Penetration | {cfm_per:,.0f} CFM |
| Total CFM (all units) | {total_cfm_raw:,.0f} CFM |
| Diversity Factor | {diversity_pct:.0f}% |
| Design CFM | {design_cfm:,.0f} CFM |
| Floor-to-Floor Height | {floor_height:.1f} ft |
| Total Shaft Height | {shaft_h:.1f} ft |
| Duct After Last Unit | {duct_after_last:.1f} ft |
| Shaft Offsets | {"Yes" if has_offset else "None"} |

### Design Criteria

| Criterion | Limit |
|---|---|
| Max floor-to-floor dP | {max_delta_p:.2f} in. WC |
| Air density (standard) | 0.075 lb/ft3 |
| Duct roughness | 0.0003 ft |
| Engineering basis | ASHRAE 2009 Duct Design Ch. 21 |

---

## PAGE 3 — RECOMMENDED SHAFT SIZE

### {pass_fail}: {label}

| Metric | Value |
|---|---|
| Shaft Dimensions | {label} |
| Gross Cross-Section | {gross_area:.0f} sq. in. |
{subduct_row}| Net Effective Area | {eff_area:.0f} sq. in. |
| Hydraulic Diameter | {dh:.1f} in. |
| Maximum Velocity | {max_vel:,.0f} FPM |

### Pressure Drop Breakdown

| Component | dP (in. WC) |
|---|---|
| Shaft Friction | {dp_shaft:.4f} |
| After-Unit Duct Loss | {dp_after:.4f} |
| Offset Losses | {dp_offset:.4f} |
| Exit / Fan Entry Loss | {dp_exit:.4f} |
| TOTAL SYSTEM dP | {dp_total:.4f} |

Maximum floor differential of {max_floor_dp:.3f} in. WC is {"within" if best.get("pass", False) else "exceeding"} the {max_delta_p:.2f} in. WC design limit.

---

## PAGE 4 — FLOOR-BY-FLOOR PRESSURE ANALYSIS

### Cumulative Airflow and Pressure Distribution

The fan on the roof pulls air upward. Floor 1 (bottom) sees zero shaft resistance. Each higher floor accumulates more friction as airflow builds.

| Floor | Cumul. CFM | Velocity (FPM) | Section dP | Accumulated dP |
|---|---|---|---|---|
{fb_rows}

The constant pressure controller compensates for varying usage by modulating the DEF fan speed in real time.

---

## PAGE 5 — FAN SELECTION: {qty_label}{fan_model}

### Selected: {qty_label}{fan_model}

| Specification | Value |
|---|---|
| Design Airflow | {design_cfm:,.0f} CFM |
| System Static Pressure | {dp_total:.4f} in. WC |
| Available CFM at Design SP | {fan_sel["available_cfm"]:,.0f} CFM |
| Capacity Margin | {margin_pct}% |
| Voltage | {fan_sel["specs"]["voltage"]} |
| Horsepower | {fan_sel["specs"]["hp"]} |
| Motor Type | {fan_sel["specs"]["motor"]} |
| Impeller | {fan_sel["specs"]["impeller"]}, 5052 Aluminum |
| Listings | ETL — UL 705; CSA C22.2 |
| Weight | {fan_sel["specs"]["weight"]} |
{"| CFM per Fan | " + f'{fan_sel["cfm_per_fan"]:,.0f} CFM |' if fan_sel.get("parallel") else ""}

### DEF Fan Family

| Model | Max CFM | Max SP | Voltage | HP | Motor | Impeller |
|---|---|---|---|---|---|---|
| DEF004 | 540 | 1.0 in. | 120V/1ph | 1/2 | EC | BI |
| DEF008 | 970 | 1.75 in. | 120V/1ph | 1/2 | EC | BI |
| DEF015 | 1860 | 2.0 in. | 120V/1ph | 1/2 | EC | BI |
| DEF025 | 2480 | 2.0 in. | 120V/1ph | 1 | EC | BI |
| DEF035 | 4100 | 4.0 in. | 208-480V/3ph | 3 | ID | BC |
| DEF050 | 5850 | 4.0 in. | 208-480V/3ph | 5 | ID | BC |

---

## PAGE 6 — CONTROLLER SELECTION: {ctrl_model}

### Selected: {ctrl_model} ({ctrl_system})

{ctrl_reason}

| Specification | Value |
|---|---|
| System | {ctrl_system} |
| Control Signal | 0-10 VDC |
| Technology | EC-Flow |
| Communication | Modbus RTU RS485 |
| Listings | {ctrl_listings} |
| Accessories | {ctrl_accessories} |

| Building Height | Controller | System | Key Feature |
|---|---|---|---|
| Up to 7 Stories | L150.H | HRS | Single pressure transducer |
| Over 7 Stories | L152.M | MES | Dual transducers, UL 864, stack effect mitigation |

---

## PAGE 7 — DEF PRODUCT DATA SHEET

### DEF Dynamic Exhaust Fan

The DEF is a variable speed exhaust fan ETL listed for Laundry Exhaust, Supply Air, APS PurgeSafe, and High-Rise Exhaust Systems. Indoor or outdoor use. Removable clean-out door. Variable speed direct drive motor.

Features: Up to 6000 CFM, up to 5 in. w.c. static pressure. G90 galvanized steel housing, 5052 Aluminum impeller. AMCA 99-0401 Type B Spark Resistant. UL 705, CSA C22.2 listed. 2-Year Warranty. Made in America.

### Capacity (1PH)

| SP | DEF004 | DEF008 | DEF015 | DEF025 |
|---|---|---|---|---|
| 0 | 540 | 970 | 1860 | 2480 |
| 0.50 | 430 | 840 | 1700 | 2320 |
| 1.00 | 240 | 680 | 1520 | 2140 |
| 1.50 | -- | 440 | 1280 | 1930 |
| 2.00 | -- | -- | 990 | 1630 |

### Capacity (3PH)

| SP | DEF035 | DEF050 |
|---|---|---|
| 0 | 4100 | 5850 |
| 1.00 | 3770 | 5450 |
| 2.00 | 3460 | 5090 |
| 3.00 | 3120 | 4680 |
| 4.00 | 2630 | 4230 |

---

## PAGE 8 — CONTROLLER DATA SHEET

### {ctrl_page_title}

{ctrl_text}

---

## PAGE 9 — CONTACT AND DISCLAIMER

### Project Summary

**Project:** {project_name}
**Location:** {project_location}
**Prepared for:** {operator_email}

### Contact LF Systems

Phone: 817-393-4029 | Web: lfsystems.net | 100 S Sylvania Ave, Fort Worth, TX 76111

### Engineering Disclaimer
This report is generated by the LF Systems MES Shaft Sizing Calculator for estimation purposes. Calculations per ASHRAE 2009 Duct Design Chapter 21. Final design must be verified by a licensed professional engineer.

**IMPORTANT:** All shaft sizing calculations and subduct area deductions are based on the use of LF Systems SD Series Low Profile Subducts. Use of alternative subduct products may result in different area deductions and require resizing of the shaft. Contact LF Systems for assistance.

2026 LF Systems, a sub-brand of RM Manifold Group Inc. All information subject to change without notice.
"""
    return text


def generate_pdf_report(ss, best, fan_sel, ctrl, chart_png_bytes) -> bytes:
    """Generate a professional PDF report with branded cover page."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table as RLTable, TableStyle, Image, PageBreak,
                                     HRFlowable, KeepTogether)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfgen import canvas as rl_canvas
    from datetime import date as _date

    # Brand colors
    NAVY = colors.HexColor('#2a3853')
    MID_BLUE = colors.HexColor('#234699')
    RED = colors.HexColor('#b11f33')
    GRAY7 = colors.HexColor('#97999b')
    GRAY5 = colors.HexColor('#c8c9c7')
    BLACK6 = colors.HexColor('#101820')
    WHITE = colors.white

    buf = io.BytesIO()

    # ── Logo path for cover page ──
    _logo_path = ""
    for _lp in [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "lf_logo.png"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "LF_Systems-Horizontal-Logo-RGB-1000__1_.png"),
        "/home/claude/lf_logo.png",
    ]:
        if os.path.exists(_lp):
            _logo_path = _lp
            break

    # Custom page builder: cover page — white bg, logo top-left, red bar bottom
    def cover_page_bg(canvas, doc):
        canvas.saveState()
        w, h = letter
        # Red accent bar at very bottom
        canvas.setFillColor(RED)
        canvas.rect(0, 0, w, 10, fill=1, stroke=0)
        # LF Systems logo top-left
        if _logo_path:
            try:
                canvas.drawImage(_logo_path, 0.6*inch, h - 1.1*inch,
                                 width=2.8*inch, height=0.7*inch,
                                 preserveAspectRatio=True, anchor='nw', mask='auto')
            except Exception:
                pass
        canvas.restoreState()

    def later_pages_bg(canvas, doc):
        """Footer for subsequent pages."""
        canvas.saveState()
        w, h = letter
        canvas.setFont('Helvetica', 7)
        canvas.setFillColor(GRAY7)
        canvas.drawString(0.75*inch, 0.4*inch,
            'LF Systems | 100 S Sylvania Ave, Fort Worth, TX 76111 | 817-393-4029 | lfsystems.net')
        canvas.drawRightString(w - 0.75*inch, 0.4*inch,
            f'Page {doc.page}')
        # Thin navy line at top
        canvas.setStrokeColor(NAVY)
        canvas.setLineWidth(2)
        canvas.line(0.75*inch, h - 0.55*inch, w - 0.75*inch, h - 0.55*inch)
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=1.3*inch, bottomMargin=0.5*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    # ── Helpers ──
    project_name = ss.project_name if hasattr(ss, 'project_name') and ss.project_name else 'Untitled Project'
    project_location = ss.project_location if hasattr(ss, 'project_location') and ss.project_location else ''
    operator_email = ss.operator_email if hasattr(ss, 'operator_email') and ss.operator_email else ''
    today_str = _date.today().strftime("%B %d, %Y")
    qty_label = f'{fan_sel["quantity"]}x ' if fan_sel["quantity"] > 1 else ''
    status = 'PASS' if best['passes'] else 'FAIL'

    # Reusable table style
    def brand_table(data, col_widths, header_color=NAVY):
        t = RLTable(data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), header_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [WHITE, colors.HexColor('#f5f5f5')]),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ]))
        return t

    # ═══════════════════════════════════════════
    # PAGE 1: COVER PAGE (white bg, logo, MES title)
    # ═══════════════════════════════════════════
    story.append(Spacer(1, 0.65*inch))

    # Main title — dark navy on white
    story.append(Paragraph('MES EXHAUST SHAFT', ParagraphStyle(
        'CoverTitle1', fontSize=34, textColor=NAVY, alignment=TA_CENTER,
        spaceAfter=4, fontName='Helvetica-Bold', leading=42)))
    story.append(Paragraph('SIZING REPORT', ParagraphStyle(
        'CoverTitle2', fontSize=34, textColor=NAVY, alignment=TA_CENTER,
        spaceAfter=10, fontName='Helvetica-Bold', leading=42)))

    # Subtitle
    story.append(Paragraph('Multi-Family Shaft Constant Pressure System', ParagraphStyle(
        'CoverSub', fontSize=13, textColor=GRAY7, alignment=TA_CENTER,
        spaceAfter=18, leading=18)))

    # Red divider
    story.append(HRFlowable(width="25%", thickness=3, color=RED,
                             spaceBefore=4, spaceAfter=22, hAlign='CENTER'))

    # Project info — generous spacing
    story.append(Paragraph(f'<b>{project_name}</b>', ParagraphStyle(
        'CoverProject', fontSize=20, textColor=BLACK6, alignment=TA_CENTER,
        spaceAfter=8, fontName='Helvetica-Bold', leading=26)))
    if project_location:
        story.append(Paragraph(project_location, ParagraphStyle(
            'CoverLoc', fontSize=14, textColor=GRAY7, alignment=TA_CENTER,
            spaceAfter=6, leading=18)))
    story.append(Paragraph(f'{ss.exhaust_type}', ParagraphStyle(
        'CoverType', fontSize=12, textColor=GRAY5, alignment=TA_CENTER,
        spaceAfter=8, leading=16)))
    if operator_email:
        story.append(Paragraph(f'Prepared by: {operator_email}', ParagraphStyle(
            'CoverEmail', fontSize=10, textColor=GRAY7, alignment=TA_CENTER,
            spaceAfter=4, leading=14)))
    story.append(Paragraph(today_str, ParagraphStyle(
        'CoverDate', fontSize=10, textColor=GRAY7, alignment=TA_CENTER,
        spaceAfter=26, leading=14)))

    # System At A Glance table — clean white style with red header
    glance_data = [
        ['SYSTEM AT A GLANCE', ''],
        ['Building Floors', str(ss.floors)],
        ['Design CFM', f'{best["design_cfm"]:,.0f} CFM'],
        ['Recommended Shaft', best['label']],
        ['Selected Fan', f'{qty_label}{fan_sel["model"]}'],
        ['Controller', ctrl['model']],
        ['System dP', f'{best["dp_total"]:.3f} in. WC'],
        ['Status', f'{status}'],
    ]
    tg = RLTable(glance_data, colWidths=[2.5*inch, 2.5*inch])
    status_color = colors.HexColor('#4CAF50') if status == 'PASS' else RED
    tg.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), RED),
        ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('SPAN', (0, 0), (-1, 0)),
        ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
        ('LEFTPADDING', (0, 0), (-1, 0), 10),
        # Alternating rows
        ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#f5f5f5')),
        ('BACKGROUND', (0, 2), (-1, 2), WHITE),
        ('BACKGROUND', (0, 3), (-1, 3), colors.HexColor('#f5f5f5')),
        ('BACKGROUND', (0, 4), (-1, 4), WHITE),
        ('BACKGROUND', (0, 5), (-1, 5), colors.HexColor('#f5f5f5')),
        ('BACKGROUND', (0, 6), (-1, 6), WHITE),
        ('BACKGROUND', (0, 7), (-1, 7), colors.HexColor('#f5f5f5')),
        ('TEXTCOLOR', (0, 1), (0, -1), BLACK6),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica'),
        ('TEXTCOLOR', (1, 1), (1, -2), BLACK6),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica-Bold'),
        ('TEXTCOLOR', (1, -1), (1, -1), status_color),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 7),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 7),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('LEFTPADDING', (0, 1), (-1, -1), 10),
        ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
        ('RIGHTPADDING', (1, 1), (1, -1), 10),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
        ('LINEBELOW', (0, 0), (-1, 0), 1, RED),
        ('LINEBELOW', (0, 1), (-1, -2), 0.5, colors.HexColor('#e8e8e8')),
    ]))
    story.append(tg)

    story.append(Spacer(1, 0.45*inch))

    # Footer branding
    story.append(Paragraph('LF Systems by RM Manifold', ParagraphStyle(
        'CoverBrand', fontSize=10, textColor=GRAY7, alignment=TA_CENTER)))

    story.append(PageBreak())

    # ═══════════════════════════════════════════
    # Styles for content pages
    # ═══════════════════════════════════════════
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'],
                               fontSize=14, textColor=NAVY, spaceBefore=16, spaceAfter=8,
                               fontName='Helvetica-Bold')
    h3_style = ParagraphStyle('H3', parent=styles['Heading3'],
                               fontSize=12, textColor=BLACK6, spaceBefore=12, spaceAfter=6)
    normal = styles['Normal']
    small = ParagraphStyle('Small', parent=normal, fontSize=8, textColor=GRAY7)

    # ═══════════════════════════════════════════
    # PAGE 2: SYSTEM SUMMARY
    # ═══════════════════════════════════════════
    story.append(Paragraph('System Summary', h2_style))
    sum_data = [
        ['Parameter', 'Value'],
        ['Exhaust Type', ss.exhaust_type],
        ['Number of Floors', str(ss.floors)],
        ['Total Penetrations', str(best['total_pens'])],
        ['Total CFM (all units)', f'{best["total_cfm"]:,.0f} CFM'],
        ['Diversity Factor', f'{ss.diversity_pct}%'],
        ['Design CFM', f'{best["design_cfm"]:,.0f} CFM'],
        ['Floor-to-Floor Height', f'{ss.floor_height} ft'],
        ['Total Shaft Height', f'{best["total_height"]} ft'],
        ['Duct After Last Unit', f'{ss.duct_after_last} ft'],
    ]
    story.append(brand_table(sum_data, [3*inch, 4*inch]))
    story.append(Spacer(1, 16))

    # ── Recommended Shaft Size ──
    story.append(Paragraph('Recommended Shaft Size', h2_style))
    shaft_data = [
        ['Parameter', 'Value'],
        ['Shaft Size', best['label']],
        ['Gross Area', f'{best["shaft_area"]} sq.in.'],
        ['Net Effective Area', f'{best["eff_area"]} sq.in.'],
        ['Hydraulic Diameter', f'{best["dh_in"]}"'],
        ['Max Velocity (top)', f'{best["velocity"]:,.0f} FPM'],
        ['Velocity Pressure', f'{best["vp"]:.4f} in. WC'],
        ['Floor Balance', f'{status} — {best["delta_p"]:.4f} in. WC (limit: {ss.max_delta_p})'],
    ]
    story.append(brand_table(shaft_data, [3*inch, 4*inch]))
    story.append(Spacer(1, 16))

    # ── Pressure Drop Breakdown ──
    story.append(Paragraph('Pressure Drop Breakdown', h2_style))
    dp_data = [
        ['Component', 'dP (in. WC)'],
        ['Shaft Friction (cumulative)', f'{best["dp_shaft"]:.4f}'],
        ['After-Unit Duct', f'{best["dp_after"]:.4f}'],
        ['Offset Losses', f'{best["dp_offset"]:.4f}'],
        ['Exit/Fan Loss', f'{best["dp_exit"]:.4f}'],
        ['TOTAL SYSTEM', f'{best["dp_total"]:.4f}'],
    ]
    t_dp = brand_table(dp_data, [4*inch, 3*inch])
    t_dp.setStyle(TableStyle([
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e8e8e8')),
    ]))
    story.append(t_dp)

    story.append(PageBreak())

    # ═══════════════════════════════════════════
    # PAGE 3: FLOOR-BY-FLOOR ANALYSIS
    # ═══════════════════════════════════════════
    story.append(Paragraph('Floor-by-Floor Pressure Analysis', h2_style))
    story.append(Paragraph(
        f'Bottom floor (Floor 1): 0.0000 in. WC — no air in shaft yet. '
        f'Top floor (Floor {ss.floors}): {best["dp_top"]:.4f} in. WC — '
        f'maximum accumulated friction. All floors within {ss.max_delta_p} in. WC limit.', normal))
    story.append(Spacer(1, 8))

    if best.get("floor_dp") and len(best["floor_dp"]) <= 30:
        fl_data = [['Floor', 'Cumul. CFM', 'Velocity (FPM)', 'Section dP', 'Accumulated dP']]
        for i in range(len(best["floor_dp"])):
            fl_data.append([
                str(i + 1),
                f'{best["section_cfm"][i]:,.0f}',
                f'{best["section_vel"][i]:,.0f}',
                f'{best["section_dp"][i]:.4f}',
                f'{best["floor_dp"][i]:.4f}',
            ])
        story.append(brand_table(fl_data,
            [0.7*inch, 1.3*inch, 1.4*inch, 1.5*inch, 1.5*inch]))

    story.append(PageBreak())

    # ═══════════════════════════════════════════
    # PAGE 4: FAN + CONTROLLER SELECTION
    # ═══════════════════════════════════════════
    story.append(Paragraph(f'Fan Selection — {qty_label}{fan_sel["model"]}', h2_style))
    fan_data = [
        ['Parameter', 'Value'],
        ['Selected Fan', f'{qty_label}{fan_sel["model"]}'],
        ['Design CFM', f'{fan_sel["design_cfm"]:,.0f} CFM'],
        ['System Static Pressure', f'{fan_sel["system_sp"]:.4f} in. WC'],
        ['Available CFM at SP', f'{fan_sel["available_cfm"]:,.0f} CFM'],
        ['Capacity Margin', f'{fan_sel["margin_pct"]}%'],
        ['Voltage', fan_sel['specs']['voltage']],
        ['HP', fan_sel['specs']['hp']],
        ['Motor Type', fan_sel['specs']['motor']],
        ['Impeller', fan_sel['specs']['impeller']],
        ['Listings', 'ETL — UL 705; CSA C22.2'],
        ['Weight', fan_sel['specs']['weight']],
    ]
    if fan_sel["parallel"]:
        fan_data.insert(3, ['CFM per Fan', f'{fan_sel["cfm_per_fan"]:,.0f} CFM'])
    story.append(brand_table(fan_data, [3*inch, 4*inch], RED))
    story.append(Spacer(1, 16))

    story.append(Paragraph(f'Controller Selection — {ctrl["model"]}', h2_style))
    ctrl_data = [
        ['Parameter', 'Value'],
        ['Controller', ctrl['model']],
        ['Description', ctrl['name']],
        ['System', ctrl['system']],
        ['Accessories', ctrl['accessories']],
        ['Selection Basis', ctrl['reason']],
        ['Listings', ctrl['listings']],
    ]
    story.append(brand_table(ctrl_data, [3*inch, 4*inch], RED))
    story.append(Spacer(1, 16))

    # ── Fan/System Curve Chart ──
    story.append(Paragraph('Fan Curve vs System Curve', h2_style))
    if chart_png_bytes:
        img_buf = io.BytesIO(chart_png_bytes)
        img = Image(img_buf, width=6.5*inch, height=3.9*inch)
        story.append(img)

    # ── Footer / Disclaimer ──
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        '<b>IMPORTANT:</b> Shaft sizing calculations and subduct area deductions are based on '
        'the use of LF Systems SD Series Low Profile Subducts. Use of alternative subduct products '
        'may result in different area deductions and require resizing of the shaft.',
        ParagraphStyle('SDDisclaimer', parent=normal, fontSize=8, textColor=RED,
                       spaceBefore=4, spaceAfter=4)))
    story.append(Paragraph(
        'Engineering calculations per ASHRAE 2009 Duct Design Chapter 21. '
        'Fan data from DEF product data sheet. '
        'This report is for estimation purposes. Final design must be verified by a licensed engineer.',
        small))

    # ═══════════════════════════════════════════
    # APPEND DATA SHEET PAGES
    # ═══════════════════════════════════════════
    ds_keys = ["DEF"]
    ctrl_model_str = ctrl.get("model", "")
    if "L152" in ctrl_model_str:
        ds_keys.append("L152")
    else:
        ds_keys.append("L150")

    for ds_key in ds_keys:
        ds_images = extract_datasheet_images(ds_key)
        if ds_images:
            story.append(PageBreak())
            story.append(Spacer(1, 2*inch))
            ds_titles = {"DEF": "DEF Series — Exhaust Fan Data Sheet",
                         "L150": "L150 Controller Data Sheet",
                         "L152": "L152 Controller Data Sheet"}
            story.append(Paragraph(
                ds_titles.get(ds_key, f"{ds_key} Data Sheet"),
                ParagraphStyle('DSTitle', parent=styles['Heading1'],
                               fontSize=22, textColor=MID_BLUE,
                               alignment=TA_CENTER, spaceAfter=12)))
            story.append(Paragraph(
                'The following pages are the official LF Systems product data sheet '
                'for reference in equipment selection and specification.',
                ParagraphStyle('DSNote', parent=styles['Normal'],
                               fontSize=11, textColor=GRAY7,
                               alignment=TA_CENTER, spaceAfter=6)))

            for pg_idx, (img_bytes, img_w, img_h) in enumerate(ds_images):
                story.append(PageBreak())
                img_buf = io.BytesIO(img_bytes)
                usable_w = 6.5 * inch
                usable_h = 8.8 * inch
                aspect = img_w / img_h
                if (usable_w / usable_h) > aspect:
                    fit_h = usable_h
                    fit_w = fit_h * aspect
                else:
                    fit_w = usable_w
                    fit_h = fit_w / aspect
                img_obj = Image(img_buf, width=fit_w, height=fit_h)
                story.append(img_obj)

    doc.build(story, onFirstPage=cover_page_bg, onLaterPages=later_pages_bg)
    buf.seek(0)
    return buf.read()


def generate_csi_spec(ss, best, fan_sel, ctrl) -> bytes:
    """Generate CSI Section 23 34 00 specification as a .docx file."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Helper
    def add_heading_text(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        return h

    def add_part(part_num, title):
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(f'PART {part_num} — {title}')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    def add_article(num, title):
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {title}')
        run.bold = True
        run.font.size = Pt(10)

    def add_item(text, indent=0):
        p = doc.add_paragraph(text, style='List Bullet')
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5 + indent * 0.25)
        pf.space_after = Pt(2)
        return p

    def add_para(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        return p

    # ── Header ──
    add_heading_text('SECTION 23 34 00', level=1)
    add_heading_text('HVAC FANS — HIGH-RISE EXHAUST SHAFT SYSTEM', level=2)

    # Project info
    project_name = ss.project_name if hasattr(ss, 'project_name') and ss.project_name else ''
    project_location = ss.project_location if hasattr(ss, 'project_location') and ss.project_location else ''
    operator_email = ss.operator_email if hasattr(ss, 'operator_email') and ss.operator_email else ''
    if project_name or project_location:
        p = doc.add_paragraph()
        if project_name:
            run = p.add_run(f'Project: {project_name}')
            run.bold = True
            run.font.size = Pt(10)
        if project_location:
            if project_name:
                p.add_run('  |  ')
            run = p.add_run(f'Location: {project_location}')
            run.font.size = Pt(10)
        if operator_email:
            p.add_run('  |  ')
            run = p.add_run(f'Prepared by: {operator_email}')
            run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # Fan model info
    qty_label = f'{fan_sel["quantity"]}x ' if fan_sel["quantity"] > 1 else ''
    fan_model = fan_sel["model"]

    # ── PART 1 — GENERAL ──
    add_part(1, 'GENERAL')

    add_article('1.01', 'SECTION INCLUDES')
    add_item(f'High-rise exhaust shaft system for {ss.exhaust_type.lower()} serving {ss.floors} floors')
    add_item(f'Variable speed exhaust fan(s): {qty_label}{fan_model}')
    add_item(f'Constant pressure controller: {ctrl["model"]}')
    add_item('Associated sensors, accessories, and wiring')
    add_item('Fire-rated exhaust shaft construction per local building code')

    add_article('1.02', 'REFERENCES')
    add_item('ASHRAE Handbook — HVAC Systems and Equipment')
    add_item('ASHRAE Handbook — Fundamentals, Chapter 21: Duct Design')
    add_item('NFPA 90A — Standard for the Installation of Air-Conditioning and Ventilating Systems')
    add_item('NFPA 90B — Standard for the Installation of Warm Air Heating and Air-Conditioning Systems')
    add_item('UL 705 — Standard for Safety Power Ventilators')
    add_item('UL 508 — Standard for Industrial Control Equipment')
    add_item('UL 864 — Standard for Control Units and Accessories for Fire Alarm Systems')
    add_item('IMC — International Mechanical Code, current edition')
    add_item('SMACNA — HVAC Duct Construction Standards')

    add_article('1.03', 'SUBMITTALS')
    add_item('Product data sheets for exhaust fan(s) and controller(s)')
    add_item('Fan performance curves showing operating point')
    add_item('System pressure drop calculations')
    add_item('Wiring diagrams')
    add_item('Installation and maintenance manuals')
    add_item('Warranty documentation')

    add_article('1.04', 'QUALITY ASSURANCE')
    add_item('Exhaust fan(s) shall be ETL listed to UL 705')
    add_item('Controller shall be UL 508 listed for industrial control equipment')
    if ctrl['model'].startswith('L152'):
        add_item('Controller shall be UL 864 listed for fire alarm system integration')
    add_item('Fan manufacturer shall provide a minimum 2-year mechanical warranty')
    add_item('System shall be designed and tested by the manufacturer')

    add_article('1.05', 'DELIVERY, STORAGE, AND HANDLING')
    add_item('Deliver products in factory-sealed packaging')
    add_item('Store in clean, dry, ventilated area protected from weather')
    add_item('Handle equipment carefully to prevent damage to impellers and motors')

    add_article('1.06', 'WARRANTY')
    add_item('Manufacturer shall provide a minimum 2-year warranty on all mechanical components')
    add_item('Controller shall carry a minimum 2-year warranty')

    # ── PART 2 — PRODUCTS ──
    add_part(2, 'PRODUCTS')

    add_article('2.01', 'MANUFACTURERS')
    add_item('LF Systems, Fort Worth, TX — www.lfsystems.net')
    add_item('No substitutions without prior written approval')

    add_article('2.02', f'EXHAUST FAN — {fan_model}')
    add_item(f'Model: {qty_label}{fan_model} Dynamic Exhaust Fan')
    add_item(f'Capacity: {best["design_cfm"]:,.0f} CFM at {best["dp_total"]:.2f} in. WC static pressure')
    add_item(f'Motor: {fan_sel["specs"]["hp"]} HP, {fan_sel["specs"]["voltage"]}, '
             f'{fan_sel["specs"]["motor"]} type, {fan_sel["specs"]["rpm"]} RPM')
    add_item(f'Impeller: {fan_sel["specs"]["impeller"]} — backward inclined (BI) or backward curved (BC), '
             f'5052 aluminum, dynamically balanced')
    add_item('Housing: G90 galvanized steel with integrated clean-out access panel')
    add_item('Construction: AMCA 99-0401 Type B Spark Resistant')
    add_item('Listings: ETL listed to UL 705; CSA C22.2')
    add_item('Suitable for indoor or outdoor installation')
    if fan_sel["parallel"]:
        add_item(f'Provide {fan_sel["quantity"]} fans piped in parallel, each rated for '
                 f'{fan_sel["cfm_per_fan"]:,.0f} CFM')

    add_article('2.03', f'CONTROLLER — {ctrl["model"]}')
    add_item(f'Model: {ctrl["model"]} — {ctrl["name"]}')
    add_item(f'System: {ctrl["system"]}')
    add_item('Function: Maintain constant negative pressure in exhaust shaft using EC-Flow Technology')
    add_item('Input: 24V AC/DC')
    add_item('Output: 0-10 VDC control signal to fan VFD or EC motor')
    add_item('Display: LCD with 4-button interface')
    add_item('Communication: Modbus RTU RS485')
    add_item(f'Listings: {ctrl["listings"]}')
    add_item(f'Accessories: {ctrl["accessories"]}')
    if ctrl['model'].startswith('L152'):
        add_item('Two pressure transducers installed above and below the neutral pressure plane '
                 'to mitigate seasonal stack effect and reverse stack effect')
        add_item('Dedicated input for smoke control system integration')

    add_article('2.04', 'ACCESSORIES')
    add_item('LP5 — Bi-directional pressure transducer')
    add_item('SLT — Silicone tubing')
    add_item('DP — Duct probe(s)')
    if 'VFD' in fan_sel["specs"].get("motor", "") or fan_sel["specs"]["motor"] == "ID":
        add_item('CFW500 Variable Frequency Drive — NEMA 4X, pre-programmed for DEF motor')

    add_article('2.05', 'EXHAUST SHAFT CONSTRUCTION')
    add_item(f'Shaft size: {best["label"]}')
    add_item(f'Gross cross-sectional area: {best["shaft_area"]} sq.in.')
    if ss.has_subducts:
        add_item(f'Net effective area (after subduct deductions): {best["eff_area"]} sq.in.')
    else:
        add_item(f'Net effective area: {best["eff_area"]} sq.in. (no subducts — full area)')
    add_item('Construction: Fire-rated shaft per local building code and NFPA requirements')
    add_item('Material: Galvanized steel ductwork, minimum 26 gauge, per SMACNA standards')
    if ss.has_subducts:
        add_item(f'Subduct penetrations: {ss.floor_data[0]["subduct_size"]}" diameter, '
                 f'{ss.floor_data[0]["penetrations"]} per floor')
        add_item('IMPORTANT: Shaft sizing calculations and subduct area deductions are based on '
                 'the use of LF Systems SD Series Low Profile Subducts. Use of alternative subduct '
                 'products may result in different area deductions and require resizing of the shaft. '
                 'No substitutions without prior written approval from LF Systems.')
    else:
        add_item(f'Wall openings: {ss.floor_data[0]["penetrations"]} per floor (no subducts in shaft)')
    add_item('Shaft shall be straight and vertical between all occupied floors')
    if ss.has_offset:
        add_item(f'Offset permitted above top floor: {ss.offset_elbows} elbow(s), '
                 f'{ss.offset_length} ft length')

    # ── PART 3 — EXECUTION ──
    add_part(3, 'EXECUTION')

    add_article('3.01', 'INSTALLATION')
    add_item('Install exhaust fan on roof or mechanical penthouse per manufacturer instructions')
    add_item('Mount fan on vibration isolators to prevent structure-borne noise')
    add_item('Install controller in accessible location per manufacturer instructions')
    add_item('Install pressure transducer(s) and duct probe(s) per manufacturer instructions')
    if ctrl['model'].startswith('L152'):
        add_item('Install upper pressure transducer above the neutral pressure plane of the building')
        add_item('Install lower pressure transducer below the neutral pressure plane of the building')
    add_item('Connect all wiring per manufacturer wiring diagrams and NEC requirements')
    add_item('Provide disconnect switch and overcurrent protection per NEC')

    add_article('3.02', 'SHAFT INSTALLATION')
    add_item('Install exhaust shaft vertically through fire-rated shaft enclosure')
    add_item('Seal all shaft penetrations with fire-rated materials per building code')
    if ss.has_subducts:
        add_item('Install subduct connections at each floor with fire/smoke dampers as required')
    else:
        add_item('Install wall openings at each floor with fire/smoke dampers as required')
    add_item('Ensure shaft is airtight — seal all joints with approved duct sealant')

    add_article('3.03', 'TESTING AND COMMISSIONING')
    add_item('Verify shaft is airtight before system startup')
    add_item('Set controller pressure setpoint per design requirements')
    add_item(f'Verify system maintains negative pressure in shaft at all floor levels')
    add_item(f'Verify differential pressure between bottom and top floors does not exceed '
             f'{ss.max_delta_p} in. WC')
    add_item(f'Verify total system airflow: {best["design_cfm"]:,.0f} CFM')
    add_item('Document all test results and provide to owner')
    add_item('Provide training to building maintenance staff on system operation')

    add_article('3.04', 'DESIGN DATA')
    add_para(f'Total CFM (all units): {best["total_cfm"]:,.0f} CFM')
    add_para(f'Design CFM ({ss.diversity_pct}% diversity): {best["design_cfm"]:,.0f} CFM')
    add_para(f'System static pressure: {best["dp_total"]:.4f} in. WC')
    add_para(f'Floor-to-floor height: {ss.floor_height} ft')
    add_para(f'Total shaft height: {best["total_height"]} ft')
    add_para(f'Maximum floor differential: {best["delta_p"]:.4f} in. WC')

    doc.add_paragraph('')
    # SD Subduct Disclaimer
    p_disc = doc.add_paragraph()
    run = p_disc.add_run('NOTE: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p_disc.add_run(
        'All shaft sizing calculations in this specification are based on the use of '
        'LF Systems SD Series Low Profile Subducts. The subduct area deductions used to '
        'determine net effective shaft area assume SD Series geometry. Use of other subduct '
        'products will change the area deductions and may require a larger shaft size. '
        'Contact LF Systems at 817-393-4029 or lfsystems.net for assistance.')
    run.font.size = Pt(9)
    run.font.italic = True

    doc.add_paragraph('')
    p = doc.add_paragraph('END OF SECTION 23 34 00')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ─────────────────────────────────────────────
# ENGINEERING FUNCTIONS
# ─────────────────────────────────────────────
def colebrook_friction_factor(dh_ft: float, velocity_fpm: float) -> float:
    """Iterative Colebrook equation for Darcy friction factor."""
    v_fps = velocity_fpm / 60.0
    Re = v_fps * dh_ft / KIN_VISC
    if Re < 1.0:
        return 0.0
    if Re < 2300:
        return 64.0 / Re
    f = 0.02  # initial guess
    for _ in range(80):
        rhs = -2.0 * math.log10(ROUGHNESS / (3.7 * dh_ft) + 2.51 / (Re * math.sqrt(f)))
        f_new = 1.0 / (rhs ** 2)
        if abs(f_new - f) < 1e-10:
            break
        f = f_new
    return f

def velocity_pressure(velocity_fpm: float) -> float:
    """Velocity pressure in inches of water column."""
    return AIR_DENSITY * (velocity_fpm / 1096.2) ** 2

def darcy_pressure_drop(length_ft: float, dh_in: float, sum_k: float, velocity_fpm: float) -> float:
    """
    Combined friction + fitting loss:
        Δp = [f*(12L/Dh_in) + ΣK] * ρ * (V/1096.2)²
    Returns in. WC.
    """
    if dh_in <= 0 or velocity_fpm <= 0:
        return 0.0
    dh_ft = dh_in / 12.0
    f = colebrook_friction_factor(dh_ft, velocity_fpm)
    friction_term = f * (length_ft / dh_ft)
    vp = velocity_pressure(velocity_fpm)
    return (friction_term + sum_k) * vp

def huebscher_equiv_diameter(a_in: float, b_in: float) -> float:
    """Circular equivalent diameter for a rectangular duct (Huebscher)."""
    return 1.30 * (a_in * b_in) ** 0.625 / (a_in + b_in) ** 0.25

def hydraulic_diameter_rect(a_in: float, b_in: float) -> float:
    """Hydraulic diameter Dh = 4A/P for rectangular duct."""
    area = a_in * b_in
    perim = 2.0 * (a_in + b_in)
    return 4.0 * area / perim

def circular_area(d_in: float) -> float:
    """Cross-sectional area of round duct in sq. inches."""
    return math.pi * (d_in / 2.0) ** 2


# ─────────────────────────────────────────────
# SHAFT SIZING ENGINE
# ─────────────────────────────────────────────
def size_shaft(params: dict) -> dict:
    """
    Main sizing calculation.
    Returns dict with best_result, alternatives list, and per-floor data.
    """
    floors           = params["floors"]
    floor_data       = params["floor_data"]       # list of dicts per floor
    floor_height     = params["floor_height"]
    duct_after_last  = params["duct_after_last"]
    diversity_pct    = params["diversity_pct"]
    max_delta_p      = params["max_delta_p"]
    shape_choice     = params["shape_choice"]
    user_diameter    = params.get("user_diameter", 0)
    user_rect_a      = params.get("user_rect_a", 0)
    user_rect_b      = params.get("user_rect_b", 0)
    offset_elbows    = params.get("offset_elbows", 0)
    offset_length    = params.get("offset_length", 0)
    offset_angle     = params.get("offset_angle", 90)

    # ── aggregate floor data ──
    total_cfm = 0
    total_pens = 0
    max_subduct_area_on_any_floor = 0
    floor_cumulative_cfm = []

    for fd in floor_data:
        pens = fd["penetrations"]
        cfm_each = fd["cfm_per_pen"]
        sub_size = fd["subduct_size"]
        total_pens += pens
        total_cfm += pens * cfm_each
        sub_area = pens * SUBDUCT_AREA.get(sub_size, 0)  # 0 if no subducts
        if sub_area > max_subduct_area_on_any_floor:
            max_subduct_area_on_any_floor = sub_area

    design_cfm = total_cfm * diversity_pct / 100.0
    total_height = floors * floor_height

    # ── offset fitting losses ──
    k_offset = 0.0
    if offset_elbows > 0:
        k_per = K_ELBOW_90 if offset_angle >= 60 else K_ELBOW_45
        k_offset = offset_elbows * k_per

    # ── CFM contributed per floor (with diversity) ──
    floor_cfm = []
    for fd in floor_data:
        floor_cfm.append(fd["penetrations"] * fd["cfm_per_pen"] * diversity_pct / 100.0)

    # ── try a specific shaft size ──
    def evaluate(shaft_area_sqin, dh_in, label, is_round, dim_a, dim_b):
        eff_area = shaft_area_sqin - max_subduct_area_on_any_floor
        if eff_area <= 0:
            return None
        eff_area_sqft = eff_area / 144.0

        # ────────────────────────────────────────────────
        # CUMULATIVE AIRFLOW MODEL — FAN ON ROOF
        # ────────────────────────────────────────────────
        # The exhaust fan is on the ROOF pulling air UPWARD.
        # Floor 1 (bottom) exhausts first — at this point
        # the shaft below has NO air, so friction is zero.
        #
        # As we go UP, each floor adds CFM to the shaft.
        # The shaft section ABOVE each floor carries the
        # cumulative CFM of that floor + all floors below.
        #
        # Shaft sections (each one floor-height tall):
        #   Section above Floor 1: carries Floor 1 CFM only
        #   Section above Floor 2: carries Floor 1+2 CFM
        #   ...
        #   Section above Floor N: carries ALL CFM
        #
        # PRESSURE AT EACH FLOOR PENETRATION:
        # The negative pressure the shaft exerts at each
        # floor's penetration is determined by what is
        # happening in the shaft AT that floor level.
        #
        # At Floor 1 (bottom): No air is in the shaft yet
        #   below this point. Air just enters. The shaft
        #   pressure at this level ≈ 0 (only minor entry
        #   effects). This floor has NO trouble exhausting.
        #
        # At Floor N (top): The shaft below is carrying
        #   cumulative CFM from ALL lower floors. The
        #   accumulated friction from all those sections
        #   has reduced the available negative pressure
        #   at this level. This floor has the MOST
        #   difficulty exhausting into the shaft.
        #
        # The ΔP we care about: the DIFFERENCE in shaft
        # pressure between Floor 1 (easiest) and Floor N
        # (hardest). This must be ≤ max_delta_p.
        # ────────────────────────────────────────────────

        # Build shaft sections bottom-to-top
        # section[i] = shaft between Floor i+1 and Floor i+2
        #              (or between Floor i+1 and the fan for the last)
        # section[i] carries cumulative CFM from floors 1..(i+1)
        section_dp = []
        section_cfm = []
        section_vel = []
        cumulative = 0.0

        for fi in range(floors):
            cumulative += floor_cfm[fi]
            vel_section = cumulative / eff_area_sqft if eff_area_sqft > 0 else 0
            dp_section = darcy_pressure_drop(floor_height, dh_in, 0, vel_section)
            section_dp.append(dp_section)
            section_cfm.append(cumulative)
            section_vel.append(vel_section)

        # Total shaft friction = sum of all section losses
        dp_shaft_total = sum(section_dp)

        # Maximum velocity (top section, full design CFM)
        vel_max = section_vel[-1] if section_vel else 0
        if vel_max < 50:
            return None

        vp_max = velocity_pressure(vel_max)

        # After-last-unit duct (carries full design CFM)
        dp_after = darcy_pressure_drop(duct_after_last, dh_in, 0, vel_max) if duct_after_last > 0 else 0.0

        # Offset section (carries full design CFM, above top floor)
        dp_offset = darcy_pressure_drop(offset_length, dh_in, k_offset, vel_max) if (offset_elbows > 0) else 0.0

        # Fan entry loss (at full velocity)
        dp_exit = K_EXIT * vp_max

        # Total system ΔP (what the fan must overcome)
        dp_total = dp_shaft_total + dp_after + dp_offset + dp_exit

        # ────────────────────────────────────────────────
        # PER-FLOOR PRESSURE (cumulative friction BELOW
        # each floor — this is how much shaft pressure
        # has been "used up" by friction by the time we
        # reach that floor level).
        #
        # Floor 1 (bottom): 0 friction below (air just
        #   entered, nothing above has happened yet from
        #   this floor's perspective looking into shaft).
        #
        # Floor 2: friction of section[0] has accumulated
        #   (section above Floor 1 carrying Floor 1 CFM).
        #
        # Floor N: friction of sections 0..(N-2) has
        #   accumulated below this point.
        #
        # This accumulated friction is what REDUCES the
        # available negative pressure at each floor.
        # ────────────────────────────────────────────────
        floor_dp_list = []

        for fi in range(floors):
            if fi == 0:
                # Floor 1: no shaft friction has accumulated
                # below this point — air is just entering
                accumulated = 0.0
            else:
                # Sum of friction in all sections BELOW this
                # floor (sections 0 through fi-1)
                accumulated = sum(section_dp[0:fi])
            floor_dp_list.append(round(accumulated, 5))

        dp_floor1 = floor_dp_list[0]     # Floor 1 (bottom) = 0
        dp_floorN = floor_dp_list[-1]    # Floor N (top) = max accumulated

        # The ΔP difference the controller must manage
        delta_p = round(dp_floorN - dp_floor1, 5)

        return {
            "label":        label,
            "is_round":     is_round,
            "dim_a":        dim_a,
            "dim_b":        dim_b,
            "shaft_area":   round(shaft_area_sqin, 2),
            "eff_area":     round(eff_area, 2),
            "dh_in":        round(dh_in, 2),
            "velocity":     round(vel_max, 0),
            "vp":           round(vp_max, 5),
            "dp_shaft":     round(dp_shaft_total, 5),
            "dp_after":     round(dp_after, 5),
            "dp_offset":    round(dp_offset, 5),
            "dp_entry":     0.0,
            "dp_exit":      round(dp_exit, 5),
            "dp_total":     round(dp_total, 5),
            "dp_bottom":    round(dp_floor1, 5),
            "dp_top":       round(dp_floorN, 5),
            "delta_p":      delta_p,
            "passes":       delta_p <= max_delta_p,
            "total_cfm":    total_cfm,
            "design_cfm":   round(design_cfm, 0),
            "total_pens":   total_pens,
            "total_height": total_height,
            "floor_dp":     floor_dp_list,
            "section_cfm":  [round(c, 0) for c in section_cfm],
            "section_vel":  [round(v, 0) for v in section_vel],
            "section_dp":   [round(d, 5) for d in section_dp],
        }

    # ── run sizing ──
    results = []

    if shape_choice in ("round_auto", "rect_auto"):
        if shape_choice == "round_auto":
            for d in ROUND_SIZES:
                area = circular_area(d)
                r = evaluate(area, d, f'{d}" Round', True, d, d)
                if r and 100 < r["velocity"] < 3000:
                    results.append(r)
        else:
            seen = set()
            for a in RECT_SIDES:
                for b in RECT_SIDES:
                    if b > a:
                        continue
                    if a / b > 4:
                        continue
                    key = (a, b)
                    if key in seen:
                        continue
                    seen.add(key)
                    area = a * b
                    dh = hydraulic_diameter_rect(a, b)
                    r = evaluate(area, dh, f'{a}" × {b}" Rect', False, a, b)
                    if r and 100 < r["velocity"] < 3000:
                        results.append(r)
        results.sort(key=lambda x: x["shaft_area"])
        best = next((r for r in results if r["passes"]), results[-1] if results else None)
        alts = [r for r in results if r["passes"]][:8]
    elif shape_choice == "round_user":
        area = circular_area(user_diameter)
        best = evaluate(area, user_diameter, f'{user_diameter}" Round', True, user_diameter, user_diameter)
        alts = [best] if best else []
    elif shape_choice == "rect_user":
        a, b = max(user_rect_a, user_rect_b), min(user_rect_a, user_rect_b)
        area = a * b
        dh = hydraulic_diameter_rect(a, b)
        best = evaluate(area, dh, f'{a}" × {b}" Rect', False, a, b)
        alts = [best] if best else []
    else:
        best = None
        alts = []

    return {"best": best, "alternatives": alts}


# ─────────────────────────────────────────────
# STREAMLIT APP
# ─────────────────────────────────────────────
def init_state():
    """Initialize session state for the chatbot."""
    defaults = {
        "step":             0,
        "messages":         [],
        "project_name":     "",
        "project_location": "",
        "operator_email":   "",
        "exhaust_type":     "",
        "floors":           0,
        "floor_data":       [],
        "same_all":         True,
        "floor_height":     0.0,
        "duct_after_last":  0.0,
        "diversity_pct":    100.0,
        "has_offset":       False,
        "has_subducts":     True,
        "offset_elbows":    0,
        "offset_length":    0.0,
        "offset_angle":     90,
        "shape_choice":     "",
        "user_diameter":    0,
        "user_rect_a":      0,
        "user_rect_b":      0,
        "max_delta_p":      0.25,
        "current_floor":    0,
        "awaiting":         "",
        "result":           None,
        "calc_done":        False,
        "editing_step":     None,  # Which step is being edited (None = not editing)
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# Map: step number → (state_key, label) for editable fields
EDITABLE_STEPS = {
    1:  ("project_name",     "Project Name"),
    2:  ("project_location", "Location"),
    3:  ("operator_email",   "Email"),
    4:  ("exhaust_type",     "Exhaust Type"),
    5:  ("has_subducts",     "Subducts"),
    6:  ("floors",           "Floors"),
    9:  ("floor_height",     "Floor Height"),
    10: ("duct_after_last",  "Duct After Last"),
    11: ("diversity_pct",    "Diversity"),
    12: ("has_offset",       "Offset"),
}


def add_msg(role: str, text: str, step: int = None):
    """Add a message with optional step association for edit tracking."""
    st.session_state.messages.append({"role": role, "text": text, "step": step})


def bot(text: str, step: int = None):
    add_msg("assistant", text, step)


def user(text: str, step: int = None):
    add_msg("user", text, step)


def reset():
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    init_state()
    st.session_state.step = 0


def edit_response(target_step: int):
    """Allow user to re-answer a specific step by rolling back to it."""
    # Remove messages from that step onward
    st.session_state.messages = [
        m for m in st.session_state.messages
        if m.get("step") is None or m["step"] < target_step
    ]
    # Reset calc if it was done
    st.session_state.calc_done = False
    st.session_state.result = None
    # Jump back to that step
    st.session_state.step = target_step
    st.session_state.editing_step = target_step


# ─── Step functions ─────────────────────────
def step_welcome():
    bot(
        "👋 **Welcome to the MES Exhaust Shaft Sizing Calculator!**\n\n"
        "This tool sizes fire-rated exhaust shafts in high-rise buildings "
        "using the **LF Systems MES** constant pressure system.\n\n"
        "**Products:** DEF · DBF · L150/L152 controllers  \n"
        "**Website:** [lfsystems.net](https://www.lfsystems.net)\n\n"
        "---\n"
        "Let's get started! First, **what is the project name?**\n\n"
        "*(e.g. \"Nakatomi Plaza\" or \"Cyberdyne Tower\")*"
    )
    st.session_state.step = 1


def process_input(user_input: str):
    """Route user input to the correct handler based on current step."""
    val = user_input.strip()
    lc = val.lower()
    step = st.session_state.step
    # Clear editing flag on new input
    st.session_state.editing_step = None

    # ─── Step 1: Project Name ───
    if step == 1:
        user(val)
        if len(val.strip()) < 2:
            bot("⚠️ Please enter a project name (at least 2 characters).")
            return
        st.session_state.project_name = val.strip()
        bot(f"✅ Project: **{st.session_state.project_name}**\n\n"
            "**What is the project location?** *(City, State — e.g. \"Dallas, TX\")*")
        st.session_state.step = 2

    # ─── Step 2: Project Location ───
    elif step == 2:
        user(val)
        if len(val.strip()) < 2:
            bot("⚠️ Please enter a city and state (e.g. \"Dallas, TX\").")
            return
        st.session_state.project_location = val.strip()
        bot(f"✅ Location: **{st.session_state.project_location}**\n\n"
            "**What is your email address?** *(for the report header)*")
        st.session_state.step = 3

    # ─── Step 3: Operator Email ───
    elif step == 3:
        user(val)
        email = val.strip()
        if "@" not in email or "." not in email:
            bot("⚠️ Please enter a valid email address (e.g. \"jsmith@company.com\").")
            return
        st.session_state.operator_email = email
        bot(f"✅ Email: **{email}**\n\n"
            "---\n"
            "Now let's size the shaft. **What type of exhaust does this shaft serve?**")
        st.session_state.step = 4

    # ─── Step 4: Exhaust type ───
    elif step == 4:
        user(val)
        if "dryer" in lc:
            st.session_state.exhaust_type = "Clothes Dryers"
        elif "bath" in lc:
            st.session_state.exhaust_type = "Bathroom Exhaust"
        elif "kitchen" in lc or "hood" in lc:
            st.session_state.exhaust_type = "Kitchen Hood Exhaust"
        else:
            bot("⚠️ Please select **Dryers**, **Bathrooms**, or **Kitchen Hoods**.")
            return
        bot(f"✅ **{st.session_state.exhaust_type}** selected.\n\n"
            "**Does this shaft have subducts (branch ducts penetrating the shaft wall)?**\n\n"
            "Select **Yes** if individual branch ducts from each unit connect into the shaft "
            "(common for dryer and bath exhaust). "
            "Select **No** if the shaft has straight wall openings with no subducts inside the shaft.")
        st.session_state.step = 5

    # ─── Step 5: Subducts yes/no ───
    elif step == 5:
        user(val)
        if lc in ("yes", "y", "true", "1"):
            st.session_state.has_subducts = True
            bot("✅ **Subducts** — shaft area will be reduced by subduct cross-sections.\n\n"
                "**How many floors have penetrations into this shaft?**")
        elif lc in ("no", "n", "false", "0"):
            st.session_state.has_subducts = False
            bot("✅ **No subducts** — full shaft area available for airflow.\n\n"
                "**How many floors have openings into this shaft?**")
        else:
            bot("⚠️ Please answer **Yes** or **No**.")
            return
        st.session_state.step = 6

    # ─── Step 6: Number of floors ───
    elif step == 6:
        user(val)
        try:
            n = int(val)
            assert 1 <= n <= 120
        except:
            bot("⚠️ Enter a number between 1 and 120.")
            return
        st.session_state.floors = n
        st.session_state.floor_data = [
            {"penetrations": 1, "subduct_size": 0, "cfm_per_pen": 0} for _ in range(n)
        ]
        same_prompt = "Are **all floors configured the same?** (same # of openings, "
        if st.session_state.has_subducts:
            same_prompt += "subduct size, "
        same_prompt += "CFM)"
        bot(f"✅ **{n} floors**.\n\n" + same_prompt)
        st.session_state.step = 7

    # ─── Step 7: Same for all? ───
    elif step == 7:
        user(val)
        pens_prompt = "**How many penetrations (openings) per floor?** (1 or 2)"
        if lc in ("yes", "y", "true", "1"):
            st.session_state.same_all = True
            bot("Great — all floors the same.\n\n" + pens_prompt)
            st.session_state.step = 8
            st.session_state.awaiting = "pens"
        else:
            st.session_state.same_all = False
            st.session_state.current_floor = 0
            bot(f"OK — per-floor config.\n\n"
                f"**Floor 1 of {st.session_state.floors}: How many penetrations?** (1 or 2)")
            st.session_state.step = 8
            st.session_state.awaiting = "pens"

    # ─── Step 8: Floor data (pens → subduct → cfm) ───
    elif step == 8:
        user(val)
        aw = st.session_state.awaiting

        if aw == "pens":
            try:
                n = int(val)
                assert n in (1, 2)
            except:
                bot("⚠️ Enter **1** or **2** penetrations per floor.")
                return
            if st.session_state.same_all:
                for fd in st.session_state.floor_data:
                    fd["penetrations"] = n
            else:
                st.session_state.floor_data[st.session_state.current_floor]["penetrations"] = n

            if st.session_state.has_subducts:
                bot(f"✅ {n} penetration(s).\n\n"
                    "**What subduct size?** (4, 6, or 8 inches)\n\n"
                    "| Size | Area Removed |\n|---|---|\n"
                    '| 4" | 15.0 sq.in. |\n| 6" | 31.5 sq.in. |\n| 8" | 54.0 sq.in. |')
                st.session_state.awaiting = "sub"
            else:
                # No subducts — skip straight to CFM
                if st.session_state.same_all:
                    for fd in st.session_state.floor_data:
                        fd["subduct_size"] = 0
                else:
                    st.session_state.floor_data[st.session_state.current_floor]["subduct_size"] = 0
                bot(f"✅ {n} opening(s) per floor (no subducts).\n\n"
                    "**How many CFM per opening?**")
                st.session_state.awaiting = "cfm"

        elif aw == "sub":
            try:
                n = int(val)
                assert n in (4, 6, 8)
            except:
                bot("⚠️ Subduct must be **4**, **6**, or **8** inches.")
                return
            if st.session_state.same_all:
                for fd in st.session_state.floor_data:
                    fd["subduct_size"] = n
            else:
                st.session_state.floor_data[st.session_state.current_floor]["subduct_size"] = n
            bot(f'✅ {n}" subduct (removes {SUBDUCT_AREA[n]} sq.in.).\n\n'
                "**How many CFM per penetration?**")
            st.session_state.awaiting = "cfm"

        elif aw == "cfm":
            try:
                c = float(val)
                assert c > 0
            except:
                bot("⚠️ Enter a CFM value greater than 0.")
                return
            if st.session_state.same_all:
                for fd in st.session_state.floor_data:
                    fd["cfm_per_pen"] = c
                bot(f"✅ {c} CFM/opening applied to all {st.session_state.floors} floors.\n\n"
                    "**What is the floor-to-floor height (ft)?**")
                st.session_state.step = 9
            else:
                st.session_state.floor_data[st.session_state.current_floor]["cfm_per_pen"] = c
                cf = st.session_state.current_floor
                if cf < st.session_state.floors - 1:
                    st.session_state.current_floor = cf + 1
                    nf = cf + 2
                    bot(f"✅ Floor {cf+1} done.\n\n"
                        f"**Floor {nf} of {st.session_state.floors}: How many penetrations?** (1 or 2)")
                    st.session_state.awaiting = "pens"
                else:
                    bot(f"✅ All {st.session_state.floors} floors configured!\n\n"
                        "**What is the floor-to-floor height (ft)?**")
                    st.session_state.step = 9

    # ─── Step 9: Floor height ───
    elif step == 9:
        user(val)
        try:
            h = float(val)
            assert h > 0
        except:
            bot("⚠️ Enter a valid height in feet.")
            return
        st.session_state.floor_height = h
        bot(f"✅ {h} ft floor-to-floor.\n\n"
            "**Length of duct from the last (top) floor penetration to the exhaust fan (ft)?**\n"
            "(Include all horizontal/vertical run after the highest connection.)")
        st.session_state.step = 10

    # ─── Step 10: Duct after last ───
    elif step == 10:
        user(val)
        try:
            d = float(val)
            assert d >= 0
        except:
            bot("⚠️ Enter 0 or a positive number of feet.")
            return
        st.session_state.duct_after_last = d
        bot(f"✅ {d} ft after last unit.\n\n"
            "**What is the diversity factor?**  \n"
            "Enter a percentage from 20 to 100.  \n"
            "(e.g., 50 = only 50% of connections active simultaneously)")
        st.session_state.step = 11

    # ─── Step 11: Diversity ───
    elif step == 11:
        user(val)
        try:
            dv = float(val.replace("%", ""))
            assert 20 <= dv <= 100
        except:
            bot("⚠️ Diversity must be between 20 and 100%.")
            return
        st.session_state.diversity_pct = dv
        bot(f"✅ {dv}% diversity.\n\n"
            "**Does the shaft offset after the last floor?**  \n"
            "(The shaft must be straight between floors, but can offset above the top floor.)")
        st.session_state.step = 12

    # ─── Step 12: Offset? ───
    elif step == 12:
        user(val)
        if lc in ("yes", "y", "true", "1"):
            st.session_state.has_offset = True
            bot("**How many elbows in the offset?** (typically 2)")
            st.session_state.step = 13
            st.session_state.awaiting = "elbows"
        else:
            st.session_state.has_offset = False
            st.session_state.offset_elbows = 0
            st.session_state.offset_length = 0
            bot("✅ No offset.\n\n"
                "**How would you like to size the shaft?**\n\n"
                "Choose one:\n"
                "- **round_auto** — find optimal round diameter\n"
                "- **rect_auto** — find optimal rectangular size\n"
                "- **round_user** — I'll specify a diameter\n"
                "- **rect_user** — I'll specify rectangular dims")
            st.session_state.step = 14

    # ─── Step 13: Offset details ───
    elif step == 13:
        user(val)
        aw = st.session_state.awaiting
        if aw == "elbows":
            try:
                n = int(val)
                assert n >= 1
            except:
                bot("⚠️ Enter number of elbows (≥ 1).")
                return
            st.session_state.offset_elbows = n
            bot(f"✅ {n} elbow(s).\n\n**Total length of the offset section (ft)?**")
            st.session_state.awaiting = "olen"
        elif aw == "olen":
            try:
                ol = float(val)
                assert ol >= 0
            except:
                bot("⚠️ Enter a length ≥ 0.")
                return
            st.session_state.offset_length = ol
            bot(f"✅ {ol} ft offset.\n\n"
                "**Elbow angle?** (Enter 45 or 90; default 90)")
            st.session_state.awaiting = "oang"
        elif aw == "oang":
            try:
                ang = int(val)
                assert ang in (45, 90)
            except:
                ang = 90
            st.session_state.offset_angle = ang
            bot(f"✅ {ang}° elbows.\n\n"
                "**How would you like to size the shaft?**\n\n"
                "- **round_auto** — find optimal round diameter\n"
                "- **rect_auto** — find optimal rectangular size\n"
                "- **round_user** — I'll specify a diameter\n"
                "- **rect_user** — I'll specify rectangular dims")
            st.session_state.step = 14

    # ─── Step 14: Shape choice ───
    elif step == 14:
        user(val)
        if lc in ("round_auto", "rect_auto", "round_user", "rect_user"):
            st.session_state.shape_choice = lc
            if lc == "round_user":
                bot("**Enter round duct diameter (inches):**")
                st.session_state.step = 15
                st.session_state.awaiting = "diam"
            elif lc == "rect_user":
                bot("**Enter rectangular dimensions as `width x height` (inches):**\n"
                    "(e.g., 24 x 18)")
                st.session_state.step = 15
                st.session_state.awaiting = "rect"
            else:
                bot("**Maximum allowable ΔP between bottom & top floors?**  \n"
                    "Max = 0.25 in. WC.  Enter your target (e.g., 0.20):")
                st.session_state.step = 16
        else:
            bot("⚠️ Choose: **round_auto**, **rect_auto**, **round_user**, or **rect_user**.")

    # ─── Step 15: User size ───
    elif step == 15:
        user(val)
        aw = st.session_state.awaiting
        if aw == "diam":
            try:
                d = float(val)
                assert 6 <= d <= 60
            except:
                bot("⚠️ Diameter must be 6–60 inches.")
                return
            st.session_state.user_diameter = d
        elif aw == "rect":
            import re
            parts = re.split(r'[x×,\s]+', val)
            try:
                a, b = float(parts[0]), float(parts[1])
                assert a >= 6 and b >= 6
            except:
                bot("⚠️ Enter two dimensions ≥ 6\", e.g. `24 x 18`.")
                return
            st.session_state.user_rect_a = max(a, b)
            st.session_state.user_rect_b = min(a, b)
        bot("**Maximum allowable ΔP between bottom & top floors?**  \n"
            "Max = 0.25 in. WC.  Enter your target:")
        st.session_state.step = 16

    # ─── Step 16: Max ΔP → run calculation ───
    elif step == 16:
        user(val)
        try:
            dp = float(val)
            assert 0.01 <= dp <= 0.25
        except:
            bot("⚠️ Enter a value between 0.01 and 0.25 in. WC.")
            return
        st.session_state.max_delta_p = dp

        # ── Build params and run ──
        params = {
            "floors":          st.session_state.floors,
            "floor_data":      st.session_state.floor_data,
            "floor_height":    st.session_state.floor_height,
            "duct_after_last": st.session_state.duct_after_last,
            "diversity_pct":   st.session_state.diversity_pct,
            "max_delta_p":     dp,
            "shape_choice":    st.session_state.shape_choice,
            "user_diameter":   st.session_state.user_diameter,
            "user_rect_a":     st.session_state.user_rect_a,
            "user_rect_b":     st.session_state.user_rect_b,
            "offset_elbows":   st.session_state.offset_elbows,
            "offset_length":   st.session_state.offset_length,
            "offset_angle":    st.session_state.offset_angle,
        }
        result = size_shaft(params)
        st.session_state.result = result
        st.session_state.calc_done = True
        st.session_state.step = 17

        if result["best"] is None:
            bot("❌ **No valid shaft size found.**\n\n"
                "The CFM may be too high or the area deductions too large for available sizes. "
                "Try adjusting your inputs.\n\nType **restart** to begin again.")
        else:
            bot("✅ **Calculation complete!** See the results below. ⬇️")

    # ─── Step 17: Post-result ───
    elif step == 17:
        user(val)
        if "restart" in lc or "new" in lc or "reset" in lc:
            reset()
            step_welcome()
        else:
            bot("Type **restart** to size another shaft.")


# ─────────────────────────────────────────────
# RENDER RESULTS
# ─────────────────────────────────────────────
def render_results():
    """Display the sizing results in a professional layout."""
    result = st.session_state.result
    if not result or not result.get("best"):
        return

    best = result["best"]
    alts = result["alternatives"]
    ss = st.session_state

    st.markdown("---")
    st.markdown(
        '<h2 style="color:#234699; margin-bottom:0; font-family:Roboto,sans-serif; font-weight:900;">📐 MES Exhaust Shaft Sizing Results</h2>',
        unsafe_allow_html=True,
    )
    st.caption(f"LF Systems MES — {ss.exhaust_type}")

    # ── System Summary ──
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### 🏗️ System Summary")
        summary = {
            "Exhaust Type":             ss.exhaust_type,
            "Number of Floors":         ss.floors,
            "Total Penetrations":       best["total_pens"],
            "Total CFM (all units)":    f'{best["total_cfm"]:,.0f} CFM',
            "Diversity Factor":         f'{ss.diversity_pct}%',
            "Design CFM (with diversity)": f'{best["design_cfm"]:,.0f} CFM',
            "Floor-to-Floor Height":    f'{ss.floor_height} ft',
            "Total Shaft Height":       f'{best["total_height"]} ft',
            "Duct After Last Unit":     f'{ss.duct_after_last} ft',
        }
        st.table(pd.DataFrame(summary.items(), columns=["Parameter", "Value"]))

    with col2:
        st.markdown("#### 📏 Recommended Shaft Size")
        shaft_info = {
            "Shaft Size":           best["label"],
            "Gross Area":           f'{best["shaft_area"]} sq.in.',
            "Net Effective Area":   f'{best["eff_area"]} sq.in.',
            "Hydraulic Diameter":   f'{best["dh_in"]}" ',
            "Max Velocity (top)":   f'{best["velocity"]:,.0f} FPM',
            "Max Velocity Pressure":f'{best["vp"]:.4f} in. WC',
        }
        st.table(pd.DataFrame(shaft_info.items(), columns=["Parameter", "Value"]))

        if best["passes"]:
            st.success(f'✅ ΔP = {best["delta_p"]:.4f} in. WC  —  **PASSES**  (≤ {ss.max_delta_p})')
        else:
            st.error(f'❌ ΔP = {best["delta_p"]:.4f} in. WC  —  **FAILS**  (> {ss.max_delta_p})')

    # ── Pressure Drop Breakdown ──
    st.markdown("#### 📊 Pressure Drop Breakdown (Full System at Max CFM)")
    dp_data = {
        "Component": ["Shaft Friction (cumulative)", "After-Unit Duct", "Offset Losses",
                       "Exit/Fan Loss", "**TOTAL SYSTEM**"],
        "ΔP (in. WC)": [
            f'{best["dp_shaft"]:.4f}',
            f'{best["dp_after"]:.4f}',
            f'{best["dp_offset"]:.4f}',
            f'{best["dp_exit"]:.4f}',
            f'**{best["dp_total"]:.4f}**',
        ],
    }
    st.table(pd.DataFrame(dp_data))

    st.markdown(
        f'**Total CFM Requirement:** {best["total_cfm"]:,.0f} CFM &nbsp;→&nbsp; '
        f'**Design CFM ({ss.diversity_pct}% diversity):** {best["design_cfm"]:,.0f} CFM'
    )

    # ── Floor Balance ──
    st.markdown("#### 🏢 Floor Balance Analysis")
    st.caption(
        "Fan is on the roof pulling air upward. At the bottom floor, no exhaust "
        "air has entered the shaft yet — accumulated friction is 0. As each floor "
        "adds air on the way up, the shaft velocity and friction increase. "
        "The top floor sees the most accumulated friction from all the air below it."
    )

    bal_data = {
        "Parameter": [
            "Bottom Floor (Floor 1) — Accumulated ΔP",
            f"Top Floor (Floor {ss.floors}) — Accumulated ΔP",
            "ΔP Difference (top − bottom)",
            f"Max Allowable (≤ {ss.max_delta_p} in. WC)",
        ],
        "Value": [
            f'{best["dp_bottom"]:.4f} in. WC',
            f'{best["dp_top"]:.4f} in. WC',
            f'{best["delta_p"]:.4f} in. WC',
            "✅ PASS" if best["passes"] else "❌ FAIL — consider larger shaft",
        ],
    }
    st.table(pd.DataFrame(bal_data))

    # ── Per-Floor Detail (expandable) ──
    if best.get("floor_dp"):
        with st.expander("📋 Per-Floor Pressure & Airflow Detail", expanded=True):
            rows = []
            for i, dp_fl in enumerate(best["floor_dp"]):
                row = {
                    "Floor": i + 1,
                    "Position": "Bottom" if i == 0 else ("Top" if i == len(best["floor_dp"])-1 else ""),
                    "Cumul. CFM Above": f'{best["section_cfm"][i]:,.0f}' if best.get("section_cfm") else "",
                    "Shaft Velocity (FPM)": f'{best["section_vel"][i]:,.0f}' if best.get("section_vel") else "",
                    "Section ΔP (in. WC)": f'{best["section_dp"][i]:.4f}' if best.get("section_dp") else "",
                    "Accumulated ΔP (in. WC)": f"{dp_fl:.4f}",
                }
                rows.append(row)
            st.table(pd.DataFrame(rows))

    # ── Alternatives ──
    if alts and len(alts) > 1:
        st.markdown("#### 🔄 Alternative Sizes (meet ΔP requirement)")
        alt_rows = []
        for a in alts:
            alt_rows.append({
                "Size":          a["label"],
                "Eff. Area (sq.in.)": a["eff_area"],
                "Velocity (FPM)":     int(a["velocity"]),
                "ΔP Diff (in. WC)":   f'{a["delta_p"]:.4f}',
                "Status":        "✅" if a["passes"] else "❌",
            })
        st.table(pd.DataFrame(alt_rows))

    # ── Fan Selection ──
    st.markdown("#### 🔧 Fan Selection — LF Systems DEF")
    fan_sel = select_fan(best["design_cfm"], best["dp_total"])
    ctrl = select_controller(ss.floors)

    qty_label = f'{fan_sel["quantity"]}x ' if fan_sel["quantity"] > 1 else ''
    fan_data = {
        "Selected Fan":         f'{qty_label}{fan_sel["model"]}',
        "Design CFM":           f'{fan_sel["design_cfm"]:,.0f} CFM',
        "System Static Pressure": f'{fan_sel["system_sp"]:.4f} in. WC',
        "Available CFM at SP":  f'{fan_sel["available_cfm"]:,.0f} CFM',
        "Capacity Margin":      f'{fan_sel["margin_pct"]}%',
        "Voltage":              fan_sel["specs"]["voltage"],
        "HP":                   fan_sel["specs"]["hp"],
        "Motor":                fan_sel["specs"]["motor"],
        "Impeller":             fan_sel["specs"]["impeller"],
    }
    if fan_sel["parallel"]:
        fan_data["CFM per Fan"] = f'{fan_sel["cfm_per_fan"]:,.0f} CFM'
    st.table(pd.DataFrame(fan_data.items(), columns=["Parameter", "Value"]))

    if fan_sel["parallel"]:
        st.warning(f'⚠️ Design CFM exceeds single fan capacity. '
                   f'**{fan_sel["quantity"]} DEF050 fans in parallel** are recommended.')

    # ── Controller Selection ──
    st.markdown("#### 🎛️ Controller Selection")
    ctrl_data = {
        "Controller":    ctrl["model"],
        "Description":   ctrl["name"],
        "System":        ctrl["system"],
        "Accessories":   ctrl["accessories"],
        "Selection Basis": ctrl["reason"],
        "Listings":      ctrl["listings"],
    }
    st.table(pd.DataFrame(ctrl_data.items(), columns=["Parameter", "Value"]))

    # ── Fan Curve vs System Curve ──
    st.markdown("#### 📈 Fan Curve vs System Curve")
    system_curve = compute_system_curve(None, best)
    chart_png = generate_fan_system_chart(fan_sel, system_curve)
    st.image(chart_png, use_container_width=True)

    st.info(
        "🌐 Visit **[lfsystems.net](https://www.lfsystems.net)** for product specifications, "
        "CAD drawings, and ordering information."
    )

    # ── Download Buttons ──
    st.markdown("---")
    st.markdown("#### 📥 Downloads")

    col_dl1, col_dl2 = st.columns(2)

    with col_dl1:
        try:
            pdf_bytes = generate_pdf_report(ss, best, fan_sel, ctrl, chart_png)
            st.download_button(
                label="📄 Download PDF Report",
                data=pdf_bytes,
                file_name=f"MES_Shaft_Sizing_{ss.project_name.replace(' ', '_') if hasattr(ss, 'project_name') and ss.project_name else 'Report'}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"PDF generation error: {e}")

    with col_dl2:
        try:
            docx_bytes = generate_csi_spec(ss, best, fan_sel, ctrl)
            st.download_button(
                label="📋 Download CSI Spec (23 34 00)",
                data=docx_bytes,
                file_name=f"CSI_23_34_00_{ss.project_name.replace(' ', '_') if hasattr(ss, 'project_name') and ss.project_name else 'MES'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"CSI spec generation error: {e}")

    # ── Engineering Notes ──
    with st.expander("📝 Engineering Notes & Methodology"):
        st.markdown("""
**Calculation Methodology:**
- **Friction Factor:** Colebrook equation (iterative) with ε = 0.0003 ft (galvanized steel)
- **Pressure Drop:** Darcy-Weisbach: `Δp = [f·(L/Dh) + ΣK] · ρ · (V/1096.2)²`
- **Rectangular Equivalence:** Huebscher equation: `De = 1.30·(a·b)^0.625 / (a+b)^0.25`
- **Hydraulic Diameter:** `Dh = 4·A / P`
- **Air Density:** 0.075 lb/ft³ (standard conditions)
- **Fan Selection:** DEF series fans, interpolated from manufacturer fan curve data
- **Controller:** L150-H for ≤7 stories, L152-M for >7 stories (stack effect mitigation)

**Fan Models:**
| Model | Max CFM | Max SP | Voltage | HP |
|---|---|---|---|---|
| DEF004 | 540 | 1.00 | 120V/1ph | 1/2 |
| DEF008 | 970 | 1.75 | 120V/1ph | 1/2 |
| DEF015 | 1,860 | 2.00 | 120V/1ph | 1/2 |
| DEF025 | 2,480 | 2.00 | 120V/1ph | 1 |
| DEF035 | 4,100 | 2.00 | 208-480V/3ph | 3 |
| DEF050 | 5,850 | 2.00 | 208-480V/3ph | 5 |

**Notes:**
- Fan is on the ROOF pulling exhaust air UP through the shaft
- Shaft must be straight between floors (no offsets between occupied floors)
- Airflow is CUMULATIVE: bottom of shaft has minimal air, top has maximum
- Bottom floor ΔP ≈ 0 (no air in shaft yet); Top floor = maximum accumulated ΔP
- The MES system maintains constant negative pressure via EC-Flow Technology™
- If CFM exceeds single DEF050 capacity, multiple DEF050s are recommended in parallel

⚠️ **IMPORTANT — SD Subduct Requirement:**
All shaft sizing calculations and subduct area deductions in this report are based on the use of
**LF Systems SD Series Low Profile Subducts**. Use of alternative subduct products may result in
different area deductions and require resizing of the exhaust shaft. No substitutions without
prior written approval from LF Systems.
        """)


# ─────────────────────────────────────────────
# MAIN APP
# ─────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="MES Shaft Sizer — LF Systems",
        page_icon="🏗️",
        layout="wide",
    )

    # ── Custom CSS — LF Systems Brand (RM Manifold Style Guide) ──
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700;900&display=swap');

    .stApp {
        font-family: 'Roboto', sans-serif;
    }

    /* Header banner — Dark Navy gradient per brand guide */
    .mes-header {
        background: linear-gradient(135deg, #2a3853 0%, #101820 100%);
        padding: 20px 28px;
        border-radius: 8px;
        margin-bottom: 20px;
        border-bottom: 4px solid #b11f33;
        display: flex;
        align-items: center;
        gap: 20px;
    }
    .mes-header-logo {
        flex-shrink: 0;
    }
    .mes-header-logo img {
        height: 48px;
        width: auto;
    }
    .mes-header-text h1 {
        color: white;
        margin: 0;
        font-size: 22px;
        font-weight: 900;
        letter-spacing: -0.3px;
        font-family: 'Roboto', sans-serif;
    }
    .mes-header-text p {
        color: #c8c9c7;
        margin: 4px 0 0 0;
        font-size: 13px;
        font-weight: 400;
    }
    .mes-header-text a {
        color: #c8c9c7;
        text-decoration: underline;
    }
    .mes-badge {
        display: inline-block;
        background: #b11f33;
        color: white;
        padding: 4px 12px;
        border-radius: 4px;
        font-weight: 900;
        font-size: 14px;
        margin-right: 10px;
        letter-spacing: 0.5px;
        font-family: 'Roboto', sans-serif;
    }

    /* Chat messages — brand aligned */
    .chat-bot {
        background: #f4f5f6;
        border-left: 3px solid #234699;
        padding: 12px 16px;
        border-radius: 2px 8px 8px 2px;
        margin: 6px 0;
        font-size: 14px;
        line-height: 1.55;
        font-family: 'Roboto', sans-serif;
        color: #101820;
    }
    .chat-user {
        background: linear-gradient(135deg, #234699, #2a3853);
        color: white;
        padding: 10px 16px;
        border-radius: 8px 8px 2px 8px;
        margin: 6px 0 6px auto;
        max-width: 70%;
        text-align: right;
        font-size: 14px;
        font-family: 'Roboto', sans-serif;
    }

    /* Tables */
    table {
        font-size: 13px !important;
        font-family: 'Roboto', sans-serif !important;
    }

    /* Streamlit overrides for brand consistency */
    .stButton > button {
        font-family: 'Roboto', sans-serif;
        font-weight: 500;
        border-radius: 4px;
    }
    h1, h2, h3, h4 {
        font-family: 'Roboto', sans-serif !important;
    }

    /* Download buttons accent */
    .stDownloadButton > button {
        background-color: #234699 !important;
        color: white !important;
        border: none !important;
        font-weight: 500;
    }
    .stDownloadButton > button:hover {
        background-color: #2a3853 !important;
    }
    /* Edit button in chat - small, subtle */
    [data-testid="stChatMessage"] button[kind="secondary"] {
        padding: 2px 8px !important;
        font-size: 0.75rem !important;
        min-height: 28px !important;
        opacity: 0.5;
    }
    [data-testid="stChatMessage"] button[kind="secondary"]:hover {
        opacity: 1.0;
        background-color: #f0f0f0 !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── Header with LF Systems Logo ──
    import base64 as _b64
    _logo_path = os.path.join(os.path.dirname(__file__), "lf_logo.png")
    _logo_b64 = ""
    if os.path.exists(_logo_path):
        with open(_logo_path, "rb") as _lf:
            _logo_b64 = _b64.b64encode(_lf.read()).decode()

    _logo_html = ""
    if _logo_b64:
        _logo_html = f'<div class="mes-header-logo"><img src="data:image/png;base64,{_logo_b64}" alt="LF Systems"></div>'

    st.markdown(f"""
    <div class="mes-header">
        {_logo_html}
        <div class="mes-header-text">
            <h1>
                <span class="mes-badge">MES</span>
                Exhaust Shaft Sizing Calculator
            </h1>
            <p>High Rise Shaft Constant Pressure System &nbsp;|&nbsp;
            DEF · DBF · L150/L152 &nbsp;|&nbsp;
            <a href="https://www.lfsystems.net" target="_blank">lfsystems.net</a></p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    init_state()

    # ── Start the conversation ──
    if st.session_state.step == 0:
        step_welcome()

    # ── Render chat history with separators and edit buttons ──
    last_was_user = False
    for idx, msg in enumerate(st.session_state.messages):
        msg_step = msg.get("step")

        if msg["role"] == "assistant":
            # Thin separator before bot message (if previous was a user answer = end of Q&A pair)
            if last_was_user and idx > 1:
                st.markdown('<hr style="border:none;border-top:1px solid #e0e0e0;margin:8px 0 12px 0;">', unsafe_allow_html=True)
            with st.chat_message("assistant", avatar="🏗️"):
                st.markdown(msg["text"])
            last_was_user = False
        else:
            with st.chat_message("user", avatar="👤"):
                col_txt, col_edit = st.columns([5, 1])
                with col_txt:
                    st.markdown(msg["text"])
                # Show edit button only for editable steps (not during active editing)
                if (msg_step is not None
                    and msg_step in EDITABLE_STEPS
                    and not st.session_state.calc_done
                    and st.session_state.editing_step is None
                    and msg_step < st.session_state.step):
                    with col_edit:
                        if st.button("✏️", key=f"edit_{idx}_{msg_step}",
                                     help=f"Change {EDITABLE_STEPS[msg_step][1]}"):
                            edit_response(msg_step)
                            st.rerun()
            last_was_user = True

    # ── Render results if calculation is done ──
    if st.session_state.calc_done and st.session_state.result:
        render_results()

    # ── Quick-select buttons ──
    step = st.session_state.step
    buttons = []
    if step == 4:
        buttons = ["Dryers", "Bathrooms", "Kitchen Hoods"]
    elif step == 5:
        buttons = ["Yes", "No"]
    elif step == 7:
        buttons = ["Yes", "No"]
    elif step == 8 and st.session_state.awaiting == "pens":
        buttons = ["1", "2"]
    elif step == 8 and st.session_state.awaiting == "sub":
        buttons = ["4", "6", "8"]
    elif step == 12:
        buttons = ["Yes", "No"]
    elif step == 14:
        buttons = ["round_auto", "rect_auto", "round_user", "rect_user"]
    elif step == 17:
        buttons = ["restart"]

    if buttons:
        cols = st.columns(len(buttons) + 2)
        for i, b in enumerate(buttons):
            if cols[i + 1].button(b, key=f"qb_{step}_{b}", use_container_width=True):
                process_input(b)
                st.rerun()

    # ── Chat input ──
    if prompt := st.chat_input("Type your answer here..."):
        process_input(prompt)
        st.rerun()

    # ── Sidebar info ──
    with st.sidebar:
        st.markdown("### 🏗️ MES System Info")
        st.markdown(
            "The **MES (Multi-Family Exhaust Shaft)** system uses a constant pressure "
            "controller to maintain a slight negative pressure in fire-rated "
            "exhaust shafts in high-rise buildings.\n\n"
            "**Applications:**\n"
            "- Clothes dryer exhaust\n"
            "- Bathroom exhaust\n"
            "- Kitchen hood exhaust\n\n"
            "**Key Components:**\n"
            "- DEF — Dryer Exhaust Fan\n"
            "- DBF — Dryer Booster Fan\n"
            "- L150/L152 — Controllers\n\n"
            "**Rules:**\n"
            "- Shaft must be straight between floors\n"
            "- Offsets only after last floor\n"
            "- 1-2 penetrations per floor\n"
            "- Subducts: 4\", 6\", or 8\"\n"
            "- Max ΔP: 0.25 in. WC\n"
            "- Diversity: 20-100%\n"
        )
        st.markdown("---")
        st.markdown(
            "🌐 **[lfsystems.net](https://www.lfsystems.net)**  \n"
            "📞 Contact your LF Systems rep for product selection."
        )
        st.markdown("---")
        if st.button("🔄 Start Over", use_container_width=True):
            reset()
            st.rerun()

        st.markdown("---")
        st.caption("v1.0 — Engineering calculations per ASHRAE 2009 Chapter 21")


if __name__ == "__main__":
    main()
